"""Simple markdown to docx with images"""
import re, os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

MD_FILE = r"C:\Users\yunqi\Desktop\03_用户使用手册_已嵌入截图.md"
DOCX_FILE = r"C:\Users\yunqi\Desktop\ruanzhu\03_用户使用手册_已嵌入截图.docx"
BASE_DIR = r"C:\Users\yunqi\Desktop"

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)

for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.18)
    s.right_margin = Cm(3.18)

with open(MD_FILE, "r", encoding="utf-8") as f:
    text = f.read()

blocks = re.split(r'\n\n+', text)
img_count = 0

for block in blocks:
    block = block.strip()
    if not block:
        continue
    lines = block.split('\n')
    first = lines[0].strip()

    # Image
    m = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', first)
    if m:
        alt = m.group(1)
        path = m.group(2)
        fp = os.path.join(BASE_DIR, path[2:]) if path.startswith('./') else os.path.join(BASE_DIR, "软著截图", os.path.basename(path))
        if os.path.exists(fp):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p.add_run().add_picture(fp, width=Inches(5.2))
                img_count += 1
            except:
                p.add_run(f"[{alt}]")
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run(alt).font.size = Pt(9)
        else:
            doc.add_paragraph(f"[图缺失: {alt}]")
        continue

    # Headings
    if first.startswith('# ') and not first.startswith('## '):
        doc.add_heading(first[2:], level=1); continue
    if first.startswith('## '):
        doc.add_heading(first[3:], level=2); continue
    if first.startswith('### '):
        doc.add_heading(first[4:], level=3); continue

    # HR
    if first in ('---', '***'):
        doc.add_paragraph('─'*50); continue

    # Blockquote
    if first.startswith('> '):
        p = doc.add_paragraph()
        r = p.add_run(first[2:])
        r.italic = True; r.font.color.rgb = RGBColor(100,100,100)
        continue

    # Table
    if first.startswith('|') and first.endswith('|'):
        rows = []
        for l in lines:
            if '---' in l: continue
            rows.append([c.strip() for c in l.strip().split('|')[1:-1]])
        if rows:
            tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
            tbl.style = 'Light Grid Accent 1'
            for ri, rc in enumerate(rows):
                for ci, ct in enumerate(rc):
                    tbl.cell(ri, ci).text = ct
            doc.add_paragraph()
        continue

    # Regular paragraph
    para = ' '.join(lines)
    p = doc.add_paragraph()
    for part in re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', para):
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9)
        else:
            p.add_run(part)

doc.save(DOCX_FILE)
print(f"[DONE] {DOCX_FILE}")
print(f"Images: {img_count} | Size: {os.path.getsize(DOCX_FILE)/1024:.1f} KB")
