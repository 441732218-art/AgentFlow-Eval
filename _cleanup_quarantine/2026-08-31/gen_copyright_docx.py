import os, re, subprocess, sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ========== 配置参数 ==========
REQUIRED_PAGES = 60
MIN_LINES_PER_PAGE = 50
FONT_SIZE_PT = 12
MARGIN_TOP_CM = 2.5
MARGIN_BOTTOM_CM = 2.5
# ==============================

# 经验校准表：(行距pt, 实测每页行数) —— 基于你的实际运行数据
# 每次运行后可以把新数据点加进来，越来越准
CALIBRATION = [
    (11.40, 54.3),   # 第1次运行: 3043行/56页
    (13.04, 46.9),   # 第2次运行: 3000行/64页
]


def clean_markdown(lines):
    cleaned = []
    in_code_block = False
    for line in lines:
        stripped = line.rstrip('\n')
        if stripped.strip().startswith('`'):
            in_code_block = not in_code_block
            continue
        if re.match(r'^#{2,6}\s+\S', stripped) and not stripped.lstrip().startswith('#!'):
            continue
        if re.match(r'^[-*_]{3,}\s*$', stripped.strip()):
            continue
        if re.match(r'^\s*</?[a-zA-Z][^>]*>\s*$', stripped):
            continue
        if re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', stripped) and '|' in stripped and '-' in stripped:
            continue
        cleaned.append(stripped)
    while cleaned and not cleaned[0].strip():
        cleaned.pop(0)
    while cleaned and not cleaned[-1].strip():
        cleaned.pop()
    return cleaned


def interpolate_line_spacing(target_lines_per_page):
    """根据校准数据线性插值，推算目标行距"""
    pts = [c[0] for c in CALIBRATION]
    lpps = [c[1] for c in CALIBRATION]

    # 按每页行数排序
    paired = sorted(zip(lpps, pts))
    lpps_sorted = [p[0] for p in paired]
    pts_sorted = [p[1] for p in paired]

    # 线性插值
    if target_lines_per_page <= lpps_sorted[0]:
        return pts_sorted[0]
    if target_lines_per_page >= lpps_sorted[-1]:
        return pts_sorted[-1]

    for i in range(len(lpps_sorted) - 1):
        if lpps_sorted[i] <= target_lines_per_page <= lpps_sorted[i + 1]:
            ratio = (target_lines_per_page - lpps_sorted[i]) / (lpps_sorted[i + 1] - lpps_sorted[i])
            return pts_sorted[i] + ratio * (pts_sorted[i + 1] - pts_sorted[i])

    return pts_sorted[0]


def get_actual_pages(docx_path):
    """
    尝试获取docx实际页数
    方法1: 用win32com (Windows + Word)
    方法2: 用LibreOffice headless
    方法3: 返回None（无法检测）
    """
    abs_path = os.path.abspath(docx_path)

    # 方法1: win32com
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(abs_path)
        pages = doc.ComputeStatistics(2)  # wdStatisticPages = 2
        doc.Close(False)
        word.Quit()
        return pages
    except Exception:
        pass

    # 方法2: LibreOffice
    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", abs_path,
             "--outdir", os.path.dirname(abs_path)],
            capture_output=True, timeout=30
        )
        if result.returncode == 0:
            pdf_path = abs_path.rsplit('.', 1)[0] + '.pdf'
            if os.path.exists(pdf_path):
                # 简单读取PDF页数
                with open(pdf_path, 'rb') as f:
                    content = f.read()
                    # 粗略统计 /Type /Page 出现次数
                    count = content.count(b'/Type /Page') - content.count(b'/Type /Pages')
                    os.remove(pdf_path)
                    if count > 0:
                        return count
    except Exception:
        pass

    return None


def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run("第 ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run_page = paragraph.add_run()
    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)
    run_page.font.name = 'Times New Roman'
    run_page.font.size = Pt(10)

    run_end = paragraph.add_run(" 页")
    run_end.font.name = 'Times New Roman'
    run_end.font.size = Pt(10)


def generate_doc(final_lines, line_spacing_pt, output_file, software_name, version):
    """生成docx文档"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_header = hp.add_run(f"{software_name} {version}")
    run_header.font.name = '宋体'
    run_header.font.size = Pt(12)
    run_header.bold = True
    run_header._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    add_page_number_footer(doc)

    for line in final_lines:
        text = line if line.strip() else " "
        p = doc.add_paragraph(text)
        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(line_spacing_pt)
        fmt.widow_control = False

        run = p.runs[0]
        run.font.name = 'Courier New'
        run.font.size = Pt(FONT_SIZE_PT)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rFonts.set(qn('w:ascii'), 'Courier New')
        rFonts.set(qn('w:hAnsi'), 'Courier New')
        rPr.append(rFonts)

    doc.save(output_file)


def generate_copyright_doc(input_file, output_file, software_name, version):
    print(f"开始处理: {input_file}")
    if not os.path.exists(input_file):
        print(f"错误: 找不到文件 '{input_file}'")
        return

    with open(input_file, 'r', encoding='utf-8') as f:
        raw_lines = f.readlines()

    print(f"原始行数: {len(raw_lines)}")
    content_lines = clean_markdown(raw_lines)
    print(f"清洗后行数: {len(content_lines)}")

    half = REQUIRED_PAGES // 2  # 30

    if len(content_lines) > REQUIRED_PAGES * MIN_LINES_PER_PAGE:
        cut_lines = MIN_LINES_PER_PAGE * half  # 1500
        front = content_lines[:cut_lines]
        back = content_lines[-cut_lines:]
        final_lines = front + back
        print(f"✅ 代码充足({len(content_lines)}行)，截取: 前{len(front)}行 + 后{len(back)}行 = {len(final_lines)}行")
    else:
        final_lines = content_lines.copy()
        print(f"✅ 全部提交: {len(final_lines)}行")

    # ===== 基于校准数据计算行距 =====
    target_lpp = len(final_lines) / REQUIRED_PAGES  # 目标每页行数
    line_spacing = interpolate_line_spacing(target_lpp)

    # 安全余量：稍微减小行距，确保不超过60页
    # （多几页比少几页好处理，但我们目标是刚好60页）
    SAFETY_MARGIN = 0.15  # pt
    line_spacing = round(line_spacing - SAFETY_MARGIN, 2)

    print(f"📐 校准推算行距: {line_spacing}pt (含{SAFETY_MARGIN}pt安全余量)")
    print(f"📐 目标每页: {target_lpp:.1f} 行")

    # ===== 生成文档 =====
    print("正在生成文档...")
    generate_doc(final_lines, line_spacing, output_file, software_name, version)

    # ===== 尝试检测实际页数并自动微调 =====
    actual_pages = get_actual_pages(output_file)

    if actual_pages is not None:
        print(f"🔍 检测到实际页数: {actual_pages} 页")

        if actual_pages != REQUIRED_PAGES:
            # 自动微调：按比例调整行距
            adjusted_spacing = round(line_spacing * actual_pages / REQUIRED_PAGES, 2)
            print(f"🔧 自动微调行距: {line_spacing}pt → {adjusted_spacing}pt")

            # 更新校准数据
            measured_lpp = len(final_lines) / actual_pages
            CALIBRATION.append((line_spacing, measured_lpp))
            print(f"📊 新增校准数据点: ({line_spacing}pt, {measured_lpp:.1f}行/页)")

            # 重新生成
            generate_doc(final_lines, adjusted_spacing, output_file, software_name, version)

            # 再次验证
            actual_pages_2 = get_actual_pages(output_file)
            if actual_pages_2 is not None:
                print(f"🔍 微调后实际页数: {actual_pages_2} 页")
                if actual_pages_2 == REQUIRED_PAGES:
                    print(f"\n✅ 完美! 正好 {REQUIRED_PAGES} 页")
                else:
                    print(f"\n⚠️ 微调后仍为 {actual_pages_2} 页，请手动检查")
            else:
                print(f"\n✅ 已微调，请打开文档确认页数")
        else:
            print(f"\n✅ 完美! 正好 {REQUIRED_PAGES} 页")
    else:
        print(f"\n⚠️ 无法自动检测页数（需要安装Word或LibreOffice）")
        print(f"   请手动打开文档确认页数")
        print(f"   如果页数不对，把实际页数告诉我，我来调整校准参数")

    print(f"\n📄 文件: {output_file}")
    print(f"   总行数: {len(final_lines)}")
    print(f"   行距: {line_spacing}pt")


if __name__ == "__main__":
    INPUT_FILE = "软著_源代码_最终版.md"
    OUTPUT_FILE = "AgentFlow-Eval智能体工作流评测平台_V1.0.0_源代码鉴别材料.docx"
    SOFTWARE_NAME = "AgentFlow-Eval 智能体工作流评测平台"
    VERSION = "V1.0"
    generate_copyright_doc(INPUT_FILE, OUTPUT_FILE, SOFTWARE_NAME, VERSION)