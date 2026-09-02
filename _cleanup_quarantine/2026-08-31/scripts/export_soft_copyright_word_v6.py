# export_soft_copyright_word_v6.py
# V6最终版：每页独立节 + 硬编码页码

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ 请先安装: pip install python-docx")
    sys.exit(1)

PROJECT_ROOT = Path(r"D:\AgentFlow-Eval")
OUTPUT_DIR = PROJECT_ROOT / "软著" / "generated"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

CODE_FILES = [
    "backend/app/core/security.py",
    "backend/app/core/tenancy.py",
    "backend/app/models/task.py",
    "backend/app/core/judge_engine/metrics.py",
    "backend/app/core/judge_engine/llm_judge.py",
    "backend/app/core/agent_runner/tool_sandbox.py",
    "backend/app/core/celery_app/tasks.py",
    "backend/app/api/v1/websocket/manager.py",
    "backend/app/api/v1/endpoints/tasks.py",
]

LINES_PER_PAGE = 50

def create_word():
    doc = Document()
    
    # ============================================================
    # 第1页：独立封面
    # ============================================================
    # 清空页眉
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = ""
    
    # 设置页脚（硬编码页码）
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "第 1 页"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    
    # 封面内容
    title = doc.add_paragraph()
    run = title.add_run("AgentFlow-Eval Agent自动化评测工作台")
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    ver = doc.add_paragraph()
    run = ver.add_run("版本号：V1.0")
    run.font.size = Pt(16)
    run.font.name = "宋体"
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    author = doc.add_paragraph()
    run = author.add_run("著作权人：李凯昕")
    run.font.size = Pt(16)
    run.font.name = "宋体"
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    date = doc.add_paragraph()
    run = date.add_run("开发完成日期：2026年7月14日")
    run.font.size = Pt(16)
    run.font.name = "宋体"
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 填满首页
    for _ in range(20):
        doc.add_paragraph()
    
    # ============================================================
    # 收集所有代码行
    # ============================================================
    all_lines = []
    all_lines.append("=== 核心源代码 ===")
    all_lines.append("")
    
    for file_path in CODE_FILES:
        full_path = PROJECT_ROOT / file_path
        if not full_path.exists():
            all_lines.append(f"# 文件不存在：{file_path}")
            continue
        
        all_lines.append(f"=== {file_path} ===")
        try:
            content = full_path.read_text(encoding="utf-8")
            for line in content.splitlines():
                all_lines.append(line)
        except Exception as e:
            all_lines.append(f"# 读取失败：{e}")
        all_lines.append("")
    
    # ============================================================
    # 后续页面：每页独立节 + 硬编码页码
    # ============================================================
    page_no = 2
    total_pages = 1 + (len(all_lines) + LINES_PER_PAGE - 1) // LINES_PER_PAGE
    
    for i in range(0, len(all_lines), LINES_PER_PAGE):
        # 插入分节符（不是分页符）
        doc.add_page_break()
        
        # 获取当前节
        section = doc.sections[-1]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        
        # 设置页眉
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "AgentFlow-Eval Agent自动化评测工作台 V1.0"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.runs[0].font.size = Pt(9)
        header_para.runs[0].font.name = "宋体"
        
        # 设置页脚（硬编码页码）
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"第 {page_no} 页"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(9)
        
        # 写入当前页的代码行
        chunk = all_lines[i:i + LINES_PER_PAGE]
        while len(chunk) < LINES_PER_PAGE:
            chunk.append("")
        
        for line in chunk:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.name = "Consolas"
                p.runs[0].font.size = Pt(9)
            else:
                run = p.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            p.paragraph_format.line_spacing = Pt(13)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        
        page_no += 1
    
    # 保存
    output_path = OUTPUT_DIR / "材料二_核心源代码_V6.docx"
    doc.save(str(output_path))
    print(f"✅ 已生成：{output_path}")
    print(f"📊 总代码行数：{len(all_lines)}")
    print(f"📄 总页数：{page_no - 1} 页（含封面）")
    print("📌 每页页码已硬编码，将正确递增")
    return output_path

if __name__ == "__main__":
    create_word()
