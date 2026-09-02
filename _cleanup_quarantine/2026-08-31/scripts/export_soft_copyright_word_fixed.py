# export_soft_copyright_word_fixed.py
# 修正版：每页严格50行 + 正确页眉

import os
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ 请先安装: pip install python-docx")
    sys.exit(1)

PROJECT_ROOT = Path(r"D:\AgentFlow-Eval")
OUTPUT_DIR = PROJECT_ROOT / "软著" / "generated"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

CODE_FILES = [
    "backend/app/core/security.py",
    "backend/app/core/tenancy.py",
    "backend/app/models/task.py",
    "backend/app/core/judge_engine/metrics.py",
    "backend/app/core/judge_engine/llm_judge.py",
    "backend/app/core/agent_runner/tool_sandbox.py",
    "backend/app/core/celery_app/tasks.py",
    "backend/app/api/v1/websocket/manager.py",
    "backend/app/api/v1/endpoints/tasks.py",
]

LINES_PER_PAGE = 50

def add_code_paragraph(doc, text):
    """安全地添加代码行，处理空行情况"""
    p = doc.add_paragraph(text)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.line_spacing = Pt(13)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p

def create_word():
    doc = Document()
    
    # 设置A4页面
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
    
    # 设置正文样式
    style = doc.styles['Normal']
    style.font.name = 'Consolas'
    style.font.size = Pt(9)
    style.paragraph_format.line_spacing = Pt(13)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    
    # ===== 页眉 =====
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "AgentFlow-Eval Agent自动化评测工作台 V1.0"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.runs[0].font.size = Pt(9)
    header_para.runs[0].font.name = "宋体"
    
    # ===== 首页封面 =====
    title = doc.add_paragraph()
    run = title.add_run("AgentFlow-Eval Agent自动化评测工作台")
    run.font.size = Pt(18)
    run.font.name = "黑体"
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    ver = doc.add_paragraph()
    run = ver.add_run("版本号：V1.0")
    run.font.size = Pt(14)
    run.font.name = "宋体"
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    author = doc.add_paragraph()
    run = author.add_run("著作权人：李凯昕")
    run.font.size = Pt(14)
    run.font.name = "宋体"
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date = doc.add_paragraph()
    run = date.add_run("开发完成日期：2026年7月14日")
    run.font.size = Pt(14)
    run.font.name = "宋体"
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 空行填满首页
    for _ in range(18):
        doc.add_paragraph()
    
    # 页码
    p = doc.add_paragraph()
    run = p.add_run("第 1 页")
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== 代码内容 =====
    doc.add_page_break()
    
    all_lines = []
    all_lines.append("=== 核心源代码 ===")
    all_lines.append("")
    
    for file_path in CODE_FILES:
        full_path = PROJECT_ROOT / file_path
        if not full_path.exists():
            all_lines.append(f"# 文件不存在：{file_path}")
            continue
        
        all_lines.append(f"=== {file_path} ===")
        try:
            content = full_path.read_text(encoding="utf-8")
            for line in content.splitlines():
                all_lines.append(line)
        except Exception as e:
            all_lines.append(f"# 读取失败：{e}")
        all_lines.append("")
    
    # 分页输出
    page_no = 2
    for i in range(0, len(all_lines), LINES_PER_PAGE):
        chunk = all_lines[i:i + LINES_PER_PAGE]
        while len(chunk) < LINES_PER_PAGE:
            chunk.append("")
        
        for line in chunk:
            add_code_paragraph(doc, line)
        
        # 页码
        p = doc.add_paragraph()
        run = p.add_run(f"第 {page_no} 页")
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        page_no += 1
        if i + LINES_PER_PAGE < len(all_lines):
            doc.add_page_break()
    
    output_path = OUTPUT_DIR / "材料二_核心源代码_修正版.docx"
    doc.save(str(output_path))
    print(f"✅ 已生成：{output_path}")
    print(f"📊 总行数：{len(all_lines)}，共 {page_no - 1} 页")
    return output_path

if __name__ == "__main__":
    create_word()
