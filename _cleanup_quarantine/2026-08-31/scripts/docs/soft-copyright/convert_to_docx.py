import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, font_cjk='宋体', font_ascii='Calibri', size=12,
                 bold=False, italic=False, color=None):
    """Set Western + East-Asian fonts via XML."""
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_ascii)
    rFonts.set(qn('w:hAnsi'), font_ascii)
    rFonts.set(qn('w:eastAsia'), font_cjk)
    rFonts.set(qn('w:cs'), font_ascii)
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_doc_default_font(doc, font_cjk='宋体', font_ascii='Calibri', size=12):
    """Set Normal style default font."""
    style = doc.styles['Normal']
    style.font.size = Pt(size)
    style.paragraph_format.line_spacing = 1.5
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_ascii)
    rFonts.set(qn('w:hAnsi'), font_ascii)
    rFonts.set(qn('w:eastAsia'), font_cjk)
    rFonts.set(qn('w:cs'), font_ascii)


def parse_inline(text, para, cjk='宋体', ascii='Calibri', sz=12):
    """Parse **bold** and `code` inline markdown."""
    rest = text
    while rest:
        bm = re.match(r'\*\*(.+?)\*\*', rest)
        cm = re.match(r'`(.+?)`', rest)
        if bm:
            r = para.add_run(bm.group(1))
            set_run_font(r, cjk, ascii, sz, bold=True)
            rest = rest[bm.end():]
        elif cm:
            r = para.add_run(cm.group(1))
            set_run_font(r, '黑体', 'Consolas', sz - 1)
            rest = rest[cm.end():]
        else:
            ns = re.search(r'\*\*|`', rest)
            if ns:
                s = rest[:ns.start()]
                if s:
                    r = para.add_run(s)
                    set_run_font(r, cjk, ascii, sz)
                rest = rest[ns.start():]
            else:
                r = para.add_run(rest)
                set_run_font(r, cjk, ascii, sz)
                break


def make_heading(doc, text, level):
    """Heading with Chinese font."""
    h = doc.add_heading(level=min(level, 3))
    r = h.add_run(text)
    sizes = {1: 18, 2: 15, 3: 13}
    set_run_font(r, '黑体', 'Arial', sizes.get(level, 13), bold=True)
    return h


def write_bullets(doc, buf):
    """Write buffered bullet items as List Bullet paragraphs."""
    for item in buf:
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        parse_inline(item, p)
    buf.clear()


def convert_md_to_docx(md_path, docx_path):
    """Convert markdown to .docx with proper Chinese font support."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    set_doc_default_font(doc)
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.18)
        sec.right_margin = Cm(3.18)

    lines = content.split('\n')
    i = 0
    buf = []

    while i < len(lines):
        ln = lines[i]

        # empty
        if not ln.strip():
            write_bullets(doc, buf)
            i += 1
            continue

        # hr
        if ln.strip() == '---':
            write_bullets(doc, buf)
            doc.add_paragraph('─' * 40)
            i += 1
            continue

        # heading
        hm = re.match(r'^(#{1,6})\s+(.+)$', ln)
        if hm:
            write_bullets(doc, buf)
            make_heading(doc, hm.group(2), len(hm.group(1)))
            i += 1
            continue

        # table
        if (ln.strip().startswith('|')
                and i + 1 < len(lines)
                and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip())):
            write_bullets(doc, buf)
            header = [c.strip() for c in ln.strip().split('|') if c.strip()]
            i += 2
            rows = [header]
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [c.strip() for c in lines[i].strip().split('|') if c.strip()]
                rows.append(row)
                i += 1
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]),
                                    style='Light Grid Accent 1')
                for ri, row in enumerate(rows):
                    for ci, ct in enumerate(row):
                        cell = tbl.cell(ri, ci)
                        cell.text = ''
                        cp = cell.paragraphs[0]
                        cr = cp.add_run(ct)
                        set_run_font(cr, size=11)
                        if ri == 0:
                            cr.bold = True
                doc.add_paragraph('')
            continue

        # blockquote
        if ln.strip().startswith('>'):
            write_bullets(doc, buf)
            qt = ln.strip()[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            r = p.add_run(qt)
            set_run_font(r, size=10, italic=True, color=(128, 128, 128))
            i += 1
            continue

        # ordered list
        om = re.match(r'^(\d+)\.\s+(.+)$', ln)
        if om:
            write_bullets(doc, buf)
            p = doc.add_paragraph()
            p.style = doc.styles['List Number']
            p.clear()
            parse_inline(om.group(2), p)
            i += 1
            continue

        # unordered list
        um = re.match(r'^\s*-\s+(.+)$', ln)
        if um:
            buf.append(um.group(1))
            i += 1
            continue

        # indented continuation
        if ln.startswith('  ') and not ln.strip().startswith('-'):
            write_bullets(doc, buf)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            parse_inline(ln.strip(), p)
            i += 1
            continue

        # plain paragraph
        write_bullets(doc, buf)
        p = doc.add_paragraph()
        parse_inline(ln.strip(), p)
        i += 1

    write_bullets(doc, buf)
    doc.save(docx_path)
    print(f'Done: {docx_path}')


if __name__ == '__main__':
    import os
    base = os.path.dirname(os.path.abspath(__file__))
    src = os.path.join(base, '全能生成材料_分册', '03_用户使用手册.md')
    dst = os.path.join(base, '全能生成材料_分册', '03_用户使用手册.docx')
    try:
        convert_md_to_docx(src, dst)
    except PermissionError:
        dst = os.path.join(base, '全能生成材料_分册', '03_用户使用手册_new.docx')
        convert_md_to_docx(src, dst)

