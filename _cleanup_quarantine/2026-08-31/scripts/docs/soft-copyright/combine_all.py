# -*- coding: utf-8 -*-
"""合并全部4份软著材料为单一Word文档"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"d:\AgentFlow-Eval\scripts\docs\soft-copyright\全能生成材料_分册"
SS   = r"d:\AgentFlow-Eval\docs\soft-copyright\screenshots"
OUT  = r"d:\AgentFlow-Eval\scripts\docs\soft-copyright\软著完整材料_AgentFlow-Eval_V1.0.docx"

SRC = [
    ("01_软件主要功能与技术特点.md", "材料一"),
    ("02_核心源代码.md",           "材料二"),
    ("03_用户使用手册.md",         "材料三"),
    ("04_软件设计说明书.md",       "材料四"),
]

def sf(run, cjk='宋体', ascii='Calibri', sz=12, b=False, i=False, c=None):
    """Set font for a run with Chinese support."""
    run.font.size = Pt(sz); run.bold = b; run.italic = i
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    for k,v in [('ascii',ascii),('hAnsi',ascii),('eastAsia',cjk),('cs',ascii)]:
        rf.set(qn(f'w:{k}'), v)
    if c: run.font.color.rgb = RGBColor(*c)

def set_doc_default(doc):
    st = doc.styles['Normal']; st.font.size = Pt(12)
    st.paragraph_format.line_spacing = 1.5
    rPr = st.element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    for k,v in [('ascii','Calibri'),('hAnsi','Calibri'),('eastAsia','宋体'),('cs','Calibri')]:
        rf.set(qn(f'w:{k}'), v)
    for sec in doc.sections:
        sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.18); sec.right_margin = Cm(3.18)

def heading(doc, text, lvl):
    h = doc.add_heading(level=min(lvl, 3))
    sz = {1:18, 2:15, 3:13}.get(lvl, 13)
    r = h.add_run(text)
    sf(r, '黑体', 'Arial', sz, True)
    return h

def add_para(doc, txt, **kw):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    sf(r, **kw)
    return p

def inline_para(doc, txt, cjk='宋体', ascii='Calibri', sz=12):
    """Paragraph with **bold**, `code` and ![image] support."""
    # Image standalone line
    im = re.match(r'^!\[(.+?)\]\((.+?)\)$', txt.strip())
    if im:
        alt,fn = im.group(1), im.group(2)
        ip = os.path.join(SS, fn)
        if os.path.exists(ip):
            pp = doc.add_paragraph()
            pp.alignment = 1  # center
            rr = pp.add_run()
            rr.add_picture(ip, width=Inches(5.2))
            cp = doc.add_paragraph()
            cp.alignment = 1
            cr = cp.add_run(f'【截图：{alt}】')
            sf(cr, '黑体', 'Arial', 9, i=True, c=(100,100,100))
        else:
            add_para(doc, f'[图未找到: {fn}]', cjk='黑体', sz=10, c=(200,0,0))
        return

    p = doc.add_paragraph()
    rest = txt
    while rest:
        bm = re.match(r'\*\*(.+?)\*\*', rest)
        cm = re.match(r'`(.+?)`', rest)
        if bm:
            r = p.add_run(bm.group(1))
            sf(r, cjk, ascii, sz, b=True)
            rest = rest[bm.end():]
        elif cm:
            r = p.add_run(cm.group(1))
            sf(r, '黑体', 'Consolas', sz-1)
            rest = rest[cm.end():]
        else:
            ns = re.search(r'\*\*|`', rest)
            if ns:
                s = rest[:ns.start()]
                if s:
                    r = p.add_run(s)
                    sf(r, cjk, ascii, sz)
                rest = rest[ns.start():]
            else:
                r = p.add_run(rest)
                sf(r, cjk, ascii, sz)
                break

def code_block(doc, code):
    for line in code.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line)
        sf(r, '黑体', 'Consolas', 8)

def add_table(doc, rows):
    if not rows: return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]), style='Light Grid Accent 1')
    for ri, row in enumerate(rows):
        for ci, ct in enumerate(row):
            c = t.cell(ri, ci); c.text = ''
            r = c.paragraphs[0].add_run(ct)
            sf(r, sz=11)
            if ri == 0: r.bold = True
    doc.add_paragraph('')

def parse_md(doc, path):
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    i, in_code, code_buf = 0, False, []

    def flush_code():
        if code_buf:
            code_block(doc, '\n'.join(code_buf))
            code_buf.clear()

    while i < len(lines):
        ln = lines[i].rstrip('\n')
        if ln.strip().startswith('```'):
            if not in_code:
                in_code = True; code_buf = []; i+=1; continue
            else:
                in_code = False; flush_code(); i+=1; continue
        if in_code:
            code_buf.append(ln); i+=1; continue

        if not ln.strip(): i+=1; continue
        if ln.strip() == '---': doc.add_paragraph('─'*40); i+=1; continue

        hm = re.match(r'^(#{1,4})\s+(.+)$', ln)
        if hm:
            heading(doc, hm.group(2), len(hm.group(1)))
            i+=1; continue

        if (ln.strip().startswith('|') and i+1<len(lines)
                and re.match(r'^\|[\s\-:|]+\|$', lines[i+1].strip())):
            header = [c.strip() for c in ln.strip().split('|') if c.strip()]
            i+=2; rows = [header]
            while i<len(lines) and lines[i].strip().startswith('|'):
                rows.append([c.strip() for c in lines[i].strip().split('|') if c.strip()])
                i+=1
            add_table(doc, rows)
            continue

        if ln.strip().startswith('>'):
            qt = ln.strip()[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            r = p.add_run(qt)
            sf(r, sz=10, i=True, c=(100,100,100))
            i+=1; continue

        inline_para(doc, ln.strip())
        i+=1

    flush_code()

def main():
    doc = Document()
    set_doc_default(doc)

    # === Cover ===
    heading(doc, 'AgentFlow-Eval Agent自动化评测工作台', 1)
    heading(doc, '软件著作权登记申请 — 全套材料', 2)
    add_para(doc, '')
    add_para(doc, '著作权人：李凯昕', sz=14, b=True)
    add_para(doc, '开发完成日期：2026年7月14日')
    add_para(doc, '版  本  号：V1.0')
    add_para(doc, '文档生成日期：2026年7月')
    doc.add_page_break()

    # === Materials ===
    names = ['材料一：软件主要功能与技术特点',
             '材料二：核心源代码（鉴别材料）',
             '材料三：用户使用手册（操作说明书）',
             '材料四：软件设计说明书（技术文档）']

    for idx, (fname, _) in enumerate(SRC):
        fpath = os.path.join(BASE, fname)
        if not os.path.exists(fpath):
            add_para(doc, f'[文件未找到: {fname}]', c=(200,0,0))
            continue
        doc.add_page_break()
        heading(doc, names[idx], 1)
        doc.add_paragraph('')
        parse_md(doc, fpath)

    doc.save(OUT)
    print(f'✅ 完成！文件已生成:')
    print(f'   {OUT}')

if __name__ == '__main__':
    main()

