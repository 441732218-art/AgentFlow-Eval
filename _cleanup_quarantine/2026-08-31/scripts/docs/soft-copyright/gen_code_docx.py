# -*- coding: utf-8 -*-
"""生成 02_核心源代码（鉴别材料）.docx — 极简可靠版"""
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD  = r"d:\AgentFlow-Eval\scripts\docs\soft-copyright\全能生成材料_分册\02_核心源代码.md"
OUT = r"D:\AgentFlow-Eval\软著\02_核心源代码（鉴别材料）.docx"

def sf(r, cjk='宋体', ascii='Calibri', sz=12, b=False):
    r.font.size = Pt(sz); r.bold = b
    rPr = r._element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    rf.set(qn('w:ascii'), ascii); rf.set(qn('w:hAnsi'), ascii)
    rf.set(qn('w:eastAsia'), cjk); rf.set(qn('w:cs'), ascii)

with open(MD, 'r', encoding='utf-8') as f:
    text = f.read()

doc = Document()
for sec in doc.sections:
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

# Header
hdr = doc.sections[0].header; hp = hdr.paragraphs[0]; hp.alignment = 1
hr = hp.add_run('核心源代码（鉴别材料） V1.0 | 著作权人：李凯昕')
sf(hr, '黑体', 'Arial', 9)

# Title page
h = doc.add_heading(level=1); r = h.add_run('核心源代码（鉴别材料）'); sf(r, '黑体', 'Arial', 16, True)
h = doc.add_heading(level=2); r = h.add_run('软件著作权登记申请 | 著作权人：李凯昕 | V1.0'); sf(r, '黑体', 'Arial', 13, True)
p = doc.add_paragraph(); r = p.add_run('开发完成日期：2026年7月14日'); sf(r, sz=14)
doc.add_page_break()

lines = text.split('\n')
i, in_code, buf = 0, False, []

while i < len(lines):
    ln = lines[i].rstrip('\n')
    if ln.strip().startswith('```'):
        if not in_code:
            in_code = True; buf = []; i += 1; continue
        else:
            in_code = False
            if buf:
                for cl in buf:
                    pp = doc.add_paragraph()
                    pp.paragraph_format.left_indent = Cm(0.3)
                    pp.paragraph_format.space_before = Pt(0)
                    pp.paragraph_format.space_after = Pt(0)
                    pp.paragraph_format.line_spacing = 1.0
                    rr = pp.add_run(cl if cl.strip() else ' ')
                    sf(rr, '黑体', 'Consolas', 8)
            buf = []; i += 1; continue
    if in_code: buf.append(ln); i += 1; continue

    if not ln.strip(): i += 1; continue
    if ln.strip() == '---': doc.add_page_break(); i += 1; continue

    hm = re.match(r'^(#{1,3})\s+(.+)$', ln)
    if hm:
        lv = len(hm.group(1)); txt = hm.group(2)
        if lv == 1: i += 1; continue  # skip main h1
        h = doc.add_heading(level=lv)
        r = h.add_run(txt); sf(r, '黑体', 'Arial', {2:13,3:12}.get(lv,12), True)
        i += 1; continue

    # plain paragraph
    pp = doc.add_paragraph()
    rest = ln.strip()
    while rest:
        bm = re.match(r'\*\*(.+?)\*\*', rest)
        if bm: rr = pp.add_run(bm.group(1)); sf(rr, b=True); rest = rest[bm.end():]
        else: rr = pp.add_run(rest); sf(rr); break
    i += 1

if buf:
    for cl in buf:
        pp = doc.add_paragraph()
        pp.paragraph_format.left_indent = Cm(0.3)
        pp.paragraph_format.space_before = Pt(0)
        pp.paragraph_format.space_after = Pt(0)
        pp.paragraph_format.line_spacing = 1.0
        rr = pp.add_run(cl if cl.strip() else ' ')
        sf(rr, '黑体', 'Consolas', 8)

doc.save(OUT)
print('OK')
