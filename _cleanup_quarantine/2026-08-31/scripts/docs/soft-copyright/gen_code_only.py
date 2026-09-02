# -*- coding: utf-8 -*-
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def sf(r, cjk='宋体', ascii='Calibri', sz=12, b=False, i=False, c=None):
    r.font.size = Pt(sz); r.bold = b; r.italic = i
    rPr = r._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    rf.set(qn('w:ascii'), ascii); rf.set(qn('w:hAnsi'), ascii)
    rf.set(qn('w:eastAsia'), cjk); rf.set(qn('w:cs'), ascii)
    if c: r.font.color.rgb = RGBColor(*c)

def make_docx():
    doc = Document()
    s = doc.styles['Normal']; s.font.size = Pt(12)
    for sec in doc.sections:
        sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    # Title
    h=doc.add_heading(level=1)
    r=h.add_run('核心源代码（鉴别材料）'); sf(r,'黑体','Arial',16,True)
    h=doc.add_heading(level=2)
    r=h.add_run('软件著作权登记申请'); sf(r,'黑体','Arial',13,True)
    p=doc.add_paragraph(); r=p.add_run('著作权人：李凯昕  版本号：V1.0'); sf(r,sz=14, b=True)
    doc.add_page_break()

    with open(r'd:\AgentFlow-Eval\scripts\docs\soft-copyright\全能生成材料_分册\02_核心源代码.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i, in_code, buf = 0, False, []
    while i < len(lines):
        ln = lines[i].rstrip('\n')

        if ln.strip().startswith('```'):
            if not in_code: in_code=True; buf=[]; i+=1; continue
            else:
                in_code=False
                for cl in buf:
                    p=doc.add_paragraph()
                    p.paragraph_format.left_indent=Cm(0.3)
                    p.paragraph_format.space_before=Pt(0)
                    p.paragraph_format.space_after=Pt(0)
                    p.paragraph_format.line_spacing=1.0
                    r=p.add_run(cl or ' ')
                    sf(r,'黑体','Consolas',8)
                buf=[]
                i+=1; continue
        if in_code: buf.append(ln); i+=1; continue

        if not ln.strip(): i+=1; continue

        if ln.strip() == '---':
            doc.add_page_break()
            i+=1; continue

        hm = re.match(r'^(#{1,4})\s+(.+)$', ln)
        if hm:
            lv = min(len(hm.group(1)), 3)
            if lv == 1: continue  # skip main title
            h=doc.add_heading(level=lv)
            r=h.add_run(hm.group(2)); sf(r,'黑体','Arial',{2:13,3:12}.get(lv,12),True)
            i+=1; continue

        # Simple paragraph
        p=doc.add_paragraph(); rest=ln.strip()
        while rest:
            bm=re.match(r'\*\*(.+?)\*\*', rest)
            if bm: r=p.add_run(bm.group(1)); sf(r,b=True); rest=rest[bm.end():]
            else:
                r=p.add_run(rest); sf(r); break
        i+=1

    if buf:
        for cl in buf:
            p=doc.add_paragraph()
            p.paragraph_format.left_indent=Cm(0.3)
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
            p.paragraph_format.line_spacing=1.0
            r=p.add_run(cl or ' '); sf(r,'黑体','Consolas',8)

    dst = r'D:\AgentFlow-Eval\软著\02_核心源代码（鉴别材料）.docx'
    doc.save(dst)
    print(f'SAVED: {dst}')

make_docx()
