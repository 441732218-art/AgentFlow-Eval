# -*- coding: utf-8 -*-
"""生成 材料二 + 材料四 的 Word 版本"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"d:\AgentFlow-Eval\scripts\docs\soft-copyright\全能生成材料_分册"
OUT  = r"D:\AgentFlow-Eval\软著"

# --- font helper ---
def sf(r, cjk='宋体', ascii='Calibri', sz=12, b=False, i=False, c=None):
    r.font.size = Pt(sz); r.bold = b; r.italic = i
    rPr = r._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    rf.set(qn('w:ascii'), ascii); rf.set(qn('w:hAnsi'), ascii)
    rf.set(qn('w:eastAsia'), cjk); rf.set(qn('w:cs'), ascii)
    if c: r.font.color.rgb = RGBColor(*c)

def setup(doc, ttl):
    s = doc.styles['Normal']; s.font.size = Pt(12)
    s.paragraph_format.line_spacing = 1.15
    rPr = s.element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts'); rPr.insert(0, rf)
    for kv in [('ascii','Calibri'),('hAnsi','Calibri'),('eastAsia','宋体'),('cs','Calibri')]:
        rf.set(qn('w:'+kv[0]), kv[1])
    for sec in doc.sections:
        sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
    h = doc.sections[0].header
    hp = h.paragraphs[0]; hp.alignment = 1
    r = hp.add_run(f'{ttl}  V1.0  |  著作权人：李凯昕')
    sf(r, '黑体', 'Arial', 9)

def hx(doc, t, lv):
    h = doc.add_heading(level=min(lv,3))
    r = h.add_run(t)
    sf(r, '黑体', 'Arial', {1:16,2:13,3:12}.get(lv,12), b=True)

def ipara(doc, txt):
    p = doc.add_paragraph(); rest = txt
    while rest:
        bm = re.match(r'\*\*(.+?)\*\*', rest)
        cm = re.match(r'`(.+?)`', rest)
        if bm: r = p.add_run(bm.group(1)); sf(r, b=True); rest = rest[bm.end():]
        elif cm: r = p.add_run(cm.group(1)); sf(r, 'Consolas','Consolas',11); rest = rest[cm.end():]
        else:
            ns = re.search(r'\*\*|`', rest)
            if ns:
                s = rest[:ns.start()]
                if s: r = p.add_run(s); sf(r)
                rest = rest[ns.start():]
            else: r = p.add_run(rest); sf(r); break

def cblock(doc, code):
    for line in code.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line or ' ')
        sf(r, '黑体', 'Consolas', 8)

def tbl(doc, rows):
    if not rows: return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]), style='Light Grid Accent 1')
    for ri, row in enumerate(rows):
        for ci, ct in enumerate(row):
            c = t.cell(ri, ci); c.text = ''
            r = c.paragraphs[0].add_run(ct); sf(r, sz=10)
            if ri == 0: r.bold = True
    doc.add_paragraph('')


def convert(md_path, docx_path, title, code_doc=False):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    setup(doc, title)

    # Title page
    hx(doc, title, 1)
    hx(doc, '软件著作权登记申请材料', 2)
    doc.add_paragraph('')
    ipara(doc, '**著作权人：** 李凯昕')
    ipara(doc, '**版本号：** V1.0')
    ipara(doc, '**开发完成日期：** 2026年7月14日')
    doc.add_page_break()

    i, in_code, buf, skip = 0, False, [], True
    while i < len(lines):
        ln = lines[i].rstrip('\n')
        if ln.strip().startswith('```'):
            if not in_code: in_code = True; buf = []; i += 1; continue
            else:
                in_code = False
                if buf: cblock(doc, '\n'.join(buf)); buf = []
                i += 1; continue
        if in_code: buf.append(ln); i += 1; continue

        if not ln.strip(): i += 1; continue
        if ln.strip() == '---':
            if code_doc: doc.add_page_break()
            else: doc.add_paragraph('─' * 40)
            i += 1; continue

        hm = re.match(r'^(#{1,4})\s+(.+)$', ln)
        if hm:
            lv, tx = len(hm.group(1)), hm.group(2)
            if skip and lv == 1: skip = False; i += 1; continue
            hx(doc, tx, lv); i += 1; continue

        if (ln.strip().startswith('|') and i+1 < len(lines)
                and re.match(r'^\|[\s\-:|]+\|$', lines[i+1].strip())):
            hdr = [c.strip() for c in ln.strip().split('|') if c.strip()]
            i += 2; rows = [hdr]
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append([c.strip() for c in lines[i].strip().split('|') if c.strip()])
                i += 1
            tbl(doc, rows); continue

        if ln.strip().startswith('>'):
            qt = ln.strip()[1:].strip()
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8)
            r = p.add_run(qt); sf(r, sz=10, i=True, c=(100,100,100))
            i += 1; continue

        ipara(doc, ln.strip()); i += 1

    if buf: cblock(doc, '\n'.join(buf))

    doc.save(docx_path)
    print(f'✅ {os.path.basename(docx_path)}')

if __name__ == '__main__':
    convert(os.path.join(BASE, '02_核心源代码.md'),
            os.path.join(OUT, '02_核心源代码（鉴别材料）.docx'),
            '核心源代码（鉴别材料）', code_doc=True)
    convert(os.path.join(BASE, '04_软件设计说明书.md'),
            os.path.join(OUT, '04_软件设计说明书.docx'),
            '软件设计说明书（技术文档）', code_doc=False)
    print('--- 全部完成 ---')
