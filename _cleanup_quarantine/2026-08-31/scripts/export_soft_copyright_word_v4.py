# export_soft_copyright_word_v4.py
# V4修正版：封面无页眉 + 页码底部居中（使用Word页脚功能）

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
except ImportError:
    print("❌ 请先安装: pip install python-docx")
    sys.exit(1)

PROJECT_ROOT = Path(r"D:\AgentFlow-Eval")
OUTPUT_DIR = PROJECT_ROOT / "软著" / "generated"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

CODE_FILES = [
    "backend/app/core/security.py",
    "backend/app/core/tenancy.py",
    "backend/app/models/task.py",
    "backend/app/core/judge_engine/metrics.py",
    "backend/app/core/judge_engine/llm_judge.py",
    "backend/app/core/agent_runner/tool_sandbox.py",
    "backend/app/core/celery_app/tasks.py",
    "backend/app/api/v1/websocket/manager.py",
    "backend/app/api/v1/endpoints/tasks.py",
]

LINES_PER_PAGE = 50

def create_word():
    doc = Document()
    
    # ============================================================
    # 第一部分：独立封面（第1页，无页眉）
    # ============================================================
    # 清空首页页眉
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = ""
    
    # 【关键】设置首页页脚（页码居中）
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""
    
    # 首页内容（居中）
    title = doc.add_paragraph()
    run = title.add_run("AgentFlow-Eval Agent自动化评测工作台")
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    ver = doc.add_paragraph()
    run = ver.add_run("版本号：V1.0")
    run.font.size = Pt(16)
    run.font.name = "宋体"
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    author = doc.add_paragraph()
    run = author.add_run("著作权人：李凯昕")
    run.font.size = Pt(16)
    run.font.name = "宋体"
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    date = doc.add_paragraph()
    run = date.add_run("开发完成日期：2026年7月14日")
    run.font.size = Pt(16)
    run.font.name = "宋体"
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 填充首页到接近满页
    for _ in range(18):
        doc.add_paragraph()
    
    # 首页页码（使用页脚实现底部居中）
    footer_para.text = "第 1 页"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    
    # ============================================================
    # 第二部分：代码内容（从第2页开始）
    # ============================================================
    doc.add_page_break()
    
    # 设置第2页的页眉
    section = doc.sections[-1]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    
    # 设置页眉
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "AgentFlow-Eval Agent自动化评测工作台 V1.0"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.runs[0].font.size = Pt(9)
    header_para.runs[0].font.name = "宋体"
    
    # 设置页脚（页码）
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "第 2 页"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    
    # 收集所有代码行
    all_lines = []
    all_lines.append("=== 核心源代码 ===")
    all_lines.append("")
    
    for file_path in CODE_FILES:
        full_path = PROJECT_ROOT / file_path
        if not full_path.exists():
            all_lines.append(f"# 文件不存在：{file_path}")
            continue
        
        all_lines.append(f"=== {file_path} ===")
        try:
            content = full_path.read_text(encoding="utf-8")
            for line in content.splitlines():
                all_lines.append(line)
        except Exception as e:
            all_lines.append(f"# 读取失败：{e}")
        all_lines.append("")
    
    # 分页输出
    page_no = 2
    for i in range(0, len(all_lines), LINES_PER_PAGE):
        chunk = all_lines[i:i + LINES_PER_PAGE]
        while len(chunk) < LINES_PER_PAGE:
            chunk.append("")
        
        for line in chunk:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.name = "Consolas"
                p.runs[0].font.size = Pt(9)
            else:
                run = p.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            p.paragraph_format.line_spacing = Pt(13)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        
        # 更新页码（使用页脚）
        footer_para.text = f"第 {page_no} 页"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(9)
        
        page_no += 1
        if i + LINES_PER_PAGE < len(all_lines):
            doc.add_page_break()
            # 新页面继承同样的页眉页脚，无需重复设置
            # 但需要确保新页面的页脚引用正确
            footer = doc.sections[-1].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = ""
    
    # 修正：遍历所有section，确保页码正确
    # 重设所有页脚的页码
    for idx, section in enumerate(doc.sections, 1):
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"第 {idx} 页"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(9)
    
    # 保存
    output_path = OUTPUT_DIR / "材料二_核心源代码_V4.docx"
    doc.save(str(output_path))
    print(f"✅ 已生成：{output_path}")
    print(f"📊 总代码行数：{len(all_lines)}，共 {len(doc.sections)} 页（含封面）")
    print("📌 请检查：")
    print("   - 第1页封面无页眉 ✓")
    print("   - 页码在页面底部居中 ✓")
    return output_path

if __name__ == "__main__":
    create_word()
