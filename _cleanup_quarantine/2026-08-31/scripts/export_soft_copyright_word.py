# export_soft_copyright_word.py
# 软著源代码Word导出工具 - 自动添加页眉

import os
import sys
from pathlib import Path
from datetime import datetime

# 尝试导入python-docx，如无则提示安装
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ 缺少 python-docx 库，请执行：pip install python-docx")
    sys.exit(1)

# 配置
PROJECT_ROOT = Path(r"D:\AgentFlow-Eval")
OUTPUT_DIR = PROJECT_ROOT / "软著" / "generated"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# 需要收录的核心代码文件
CODE_FILES = [
    # 模块A：API Key鉴权
    "backend/app/core/security.py",
    # 模块B：轻量多租户
    "backend/app/core/tenancy.py",
    # 模块C：任务状态机与领域模型
    "backend/app/models/task.py",
    # 模块D：规则指标
    "backend/app/core/judge_engine/metrics.py",
    # 模块E：混合评分引擎
    "backend/app/core/judge_engine/llm_judge.py",
    # 模块F：工具沙箱
    "backend/app/core/agent_runner/tool_sandbox.py",
    # 模块G：异步评测编排
    "backend/app/core/celery_app/tasks.py",
    # 模块H：WebSocket推送
    "backend/app/api/v1/websocket/manager.py",
    # 模块I：评测任务列表API
    "backend/app/api/v1/endpoints/tasks.py",
]

def create_word_with_header():
    """生成带页眉的Word文档"""
    print("📄 正在生成Word文档...")
    doc = Document()
    
    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)   # 装订侧
        section.right_margin = Cm(2.0)
    
    # 添加页眉
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "AgentFlow-Eval Agent自动化评测工作台 V1.0"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.runs[0].font.size = Pt(9)
    header_para.runs[0].font.name = "宋体"
    
    # 封面标题
    title = doc.add_heading("AgentFlow-Eval 核心源代码", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"版本号：V1.0")
    doc.add_paragraph(f"著作权人：李凯昕")
    doc.add_paragraph(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
    doc.add_paragraph("=" * 50)
    
    total_lines = 0
    
    # 逐个添加代码文件
    for idx, file_path in enumerate(CODE_FILES, 1):
        full_path = PROJECT_ROOT / file_path
        if not full_path.exists():
            doc.add_paragraph(f"⚠️ 文件不存在：{file_path}")
            print(f"   ⚠️ 跳过不存在的文件：{file_path}")
            continue
        
        print(f"   [{idx}/{len(CODE_FILES)}] 正在添加：{file_path}")
        
        # 文件分隔标题
        doc.add_heading(f"=== {file_path} ===", level=1)
        
        # 读取并写入代码
        try:
            content = full_path.read_text(encoding="utf-8")
            lines = content.splitlines()
            file_line_count = 0
            
            for line in lines:
                if line.strip() or file_line_count == 0:  # 保留空行，但不过度
                    p = doc.add_paragraph(line)
                    p.runs[0].font.name = "Consolas"
                    p.runs[0].font.size = Pt(9)
                    p.paragraph_format.line_spacing = Pt(13)
                    file_line_count += 1
            
            total_lines += file_line_count
            print(f"      添加了 {file_line_count} 行代码")
            
        except Exception as e:
            doc.add_paragraph(f"读取失败：{e}")
            print(f"   ❌ 读取失败：{e}")
        
        # 文件之间加分页（最后一个不加）
        if idx < len(CODE_FILES):
            doc.add_page_break()
    
    # 保存
    output_path = OUTPUT_DIR / "材料二_核心源代码.docx"
    doc.save(str(output_path))
    
    print(f"\n✅ Word文档已生成：{output_path}")
    print(f"📊 总计收录代码行数：{total_lines}")
    print(f"📁 文件大小：{output_path.stat().st_size / 1024:.2f} KB")
    return output_path

if __name__ == "__main__":
    create_word_with_header()
