# gen_manual_docx.py — AgentFlow-Eval 软件说明书生成器（带截图）
# 用法: python gen_manual_docx.py
import re, os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SOFTWARE_NAME = "AgentFlow-Eval 智能体工作流评测平台"
VERSION = "V1.0"
OUTPUT_FILE = "AgentFlow-Eval智能体工作流评测平台_V1.0.0_软件说明书.docx"
SCREENSHOT_DIR = r"C:\Users\yunqi\Desktop\软著截图"
IMAGE_WIDTH = Inches(5.0)

# ========== 截图 → 章节映射 ==========
SCREENSHOT_MAP = {
    "4.1 系统登录":       [("screenshot_01_login.png",      "图 4-1 系统登录页面")],
    "5.1 任务列表":       [("screenshot_02_tasklist.png",   "图 5-1 评测任务列表")],
    "5.2 创建评测任务":   [("screenshot_03_create_step1.png","图 5-2 创建评测任务—基本信息配置"),
                          ("screenshot_04_create_step2.png","图 5-3 创建评测任务—数据集与模型选择"),
                          ("screenshot_05_create_step3.png","图 5-4 创建评测任务—指标与输出配置")],
    "5.3 测试用例管理":   [("screenshot_06_testcase.png",   "图 5-5 测试用例管理")],
    "5.4 任务创建总览":   [("task_create.png",              "图 5-6 任务创建流程总览")],
    "6.1 基准评测":       [("benchmark.png",                "图 6-1 基准评测配置与结果")],
    "6.2 A/B 实验":       [("ab_experiment.png",            "图 6-2 A/B 实验对比")],
    "7.1 插件管理":       [("plugins.png",                  "图 7-1 插件管理")],
    "8.1 计费管理":       [("billing.png",                  "图 8-1 计费管理")],
    "9.1 系统设置":       [("settings.png",                 "图 9-1 系统设置")],
}

# ========== 说明书正文 ==========
MANUAL_CONTENT = r"""
# 第一章 引言

## 1.1 编写目的

本手册为 AgentFlow-Eval 智能体工作流评测平台 V1.0（以下简称"本平台"）编写，旨在帮助用户全面了解本平台的功能模块，掌握平台的安装部署与操作方法，充分利用平台开展智能体工作流的评测、分析与审计工作。

本手册适用于软件测试人员、算法工程师、项目管理人员及系统运维人员使用。读者应具备基本的计算机操作能力，并了解大语言模型与智能体的基本概念。

## 1.2 软件概述

AgentFlow-Eval 智能体工作流评测平台是面向智能体（Agent）工作流的综合性评测系统。平台集成了数据集管理、模型管理、评测任务编排、指标计算、报告生成与日志审计等能力，为用户提供一站式评测服务。

平台采用模块化设计，主要包含评测任务管理、测试用例管理、基准评测、A/B 实验、插件管理、计费管理、系统设置等模块，各模块之间通过统一的接口与流程协作，保证评测数据的一致性与可追溯性。

平台主要特点包括：支持多种数据格式的上传与校验；支持多类型模型的注册与管理；支持评测工作流的灵活编排；内置多种常用评测指标；支持评测报告的一键导出；支持运行日志的自动归档与审计；支持 A/B 实验对比分析；支持插件化扩展。

## 1.3 适用范围

本手册适用于 AgentFlow-Eval 智能体工作流评测平台 V1.0，覆盖平台的安装、部署、使用与维护全过程。若软件版本升级导致界面或功能调整，以实际软件为准，本手册将同步更新。

## 1.4 术语与缩写

智能体（Agent）：指能够自主感知环境、做出决策并执行任务的软件实体。

工作流（Workflow）：指由若干节点按预定义逻辑组合而成、用于完成复杂任务的执行流程。

评测任务：指在平台上发起的一次完整评测活动，包含评测对象、数据集、指标与执行参数。

数据集：指用于评测的测试样本集合，包含输入、参考答案与约束条件。

指标：指用于衡量评测结果的量化标准，如准确率、F1 值、任务成功率等。

A/B 实验：指同时对两个或多个模型版本进行评测并对比结果的实验方法。

基准评测（Benchmark）：指基于标准数据集对模型进行的标准化评测。

# 第二章 运行环境

## 2.1 硬件环境

部署本平台前，请确认服务器硬件满足以下要求。涉及大规模数据集或高并发评测的场景，请按推荐配置进行准备。

表 2-1 硬件环境要求

| 部件 | 最低配置 | 推荐配置 |
| CPU | 4 核 2.4GHz | 8 核 3.0GHz 及以上 |
| 内存 | 16GB | 32GB 及以上 |
| 磁盘 | 100GB 可用空间 | 500GB SSD 及以上 |
| GPU | 不要求 | 16GB 显存及以上（可选） |
| 网络 | 千兆网卡 | 千兆网卡 |

## 2.2 软件环境

平台运行所需的软件环境如下表所示，请提前完成安装与配置。

表 2-2 软件环境要求

| 类别 | 名称 | 版本要求 |
| 操作系统 | Ubuntu / CentOS / Windows Server | 20.04+ / 7.9+ / 2019+ |
| 运行环境 | Python | 3.9 及以上 |
| 数据库 | PostgreSQL / MySQL | 12+ / 8.0+ |
| 缓存 | Redis | 6.0 及以上 |
| 浏览器 | Chrome / Edge | 90 及以上 |

# 第三章 安装与部署

## 3.1 安装准备

安装前请从官方渠道获取安装包，并校验文件完整性。请确认当前用户具备安装目录的读写权限，数据库账号已创建并授权。

1. 登录服务器，创建安装目录 /opt/agentflow-eval。
2. 上传安装包 agentflow-eval-v1.0.0.tar.gz 至安装目录。
3. 使用 md5sum 命令校验安装包完整性，并与官方提供的校验值比对。
4. 执行 tar -zxvf agentflow-eval-v1.0.0.tar.gz 解压安装包。

## 3.2 配置文件说明

解压完成后，进入 conf 目录编辑配置文件 config.yaml。主要配置项包括：数据库连接地址与账号、Redis 地址、服务端口、日志目录、日志归档保留天数等。

其中日志归档保留天数 retention_days 默认值为 30，表示系统每日自动归档超过 30 天的历史日志，可根据磁盘容量调整。

## 3.3 服务启动与验证

1. 执行 sh bin/init_db.sh 初始化数据库表结构。
2. 执行 sh bin/start.sh 启动平台服务。
3. 执行 sh bin/status.sh 查看各服务运行状态，确认全部为 running。
4. 打开浏览器访问 http://服务器IP:8080，看到登录页面即表示部署成功。

## 3.4 停止服务与卸载

1. 执行 sh bin/stop.sh 停止平台服务。
2. 备份数据库与日志归档目录中的重要数据。
3. 删除安装目录，完成卸载。

# 第四章 登录与主界面

## 4.1 系统登录

1. 打开浏览器，输入平台访问地址。
2. 输入用户名与密码（初始账号密码见交付清单）。
3. 点击"登录"按钮进入系统。

连续 3 次输入错误密码，账号将被锁定 15 分钟。如需紧急解锁，请联系系统管理员。

## 4.2 主界面概述

登录成功后进入主界面。主界面顶部为全局导航栏，展示软件名称、版本号与用户信息；左侧为功能菜单，包含评测任务、测试用例、基准评测、A/B 实验、插件管理、计费管理、系统设置等入口；中部为工作区，用于展示列表、表单与详情；底部为状态栏，展示当前登录用户与系统时间。

## 4.3 个人设置与退出

点击导航栏右侧用户名，选择"修改密码"，输入原密码与新密码后提交即可。修改密码后需重新登录。选择"退出登录"可安全退出系统，系统将清除本地会话信息。

# 第五章 评测任务管理

## 5.1 任务列表

进入"评测任务"页面，列表展示任务名称、关联模型、关联数据集、状态、创建时间等信息。任务状态包括：未开始、运行中、已成功、已失败、已归档。支持按名称搜索、按状态筛选、按时间排序，便于快速定位目标任务。

## 5.2 创建评测任务

创建评测任务分为三个步骤：基本信息配置、数据集与模型选择、指标与输出配置。

步骤一：填写任务名称、描述、优先级等基本信息。

步骤二：选择关联的数据集与待评测模型，配置模型调用参数。

步骤三：选择评测指标，配置输出目录与报告格式。评测配置必须包含模型名称、数据集路径、指标、输出目录四个必要字段，缺少任一字段时系统提示"评测配置缺少必要字段"并阻止保存。校验通过后提示"评测配置校验通过"。

## 5.3 测试用例管理

在"测试用例"页面可查看和管理评测所用的测试用例。支持按标签筛选、按难度分级、批量导入导出。每个测试用例包含输入、预期输出、约束条件和评分规则。

## 5.4 任务创建总览

任务创建完成后，系统自动生成任务摘要卡片，展示关键配置信息与预估耗时。用户可在总览页面一键启动任务或返回修改配置。

## 5.5 任务执行与监控

在任务列表中点击"运行"，任务进入运行中状态。监控页面实时展示执行进度、各节点状态与日志输出。任务完成后系统自动计算所选指标并生成评测报告。

## 5.6 任务重试与终止

失败任务支持"重试"，系统仅重跑失败节点及其下游节点，提高排错效率。运行中的任务可点击"终止"强制停止，终止操作不可撤销，已完成节点的结果保留。

# 第六章 评测分析

## 6.1 基准评测

基准评测模块提供基于标准数据集的标准化评测能力。用户可选择内置基准数据集或上传自定义数据集，系统自动执行评测并生成包含准确率、F1 值、BLEU、ROUGE 等指标的评测报告。

## 6.2 A/B 实验

A/B 实验模块支持同时对两个或多个模型版本进行评测并对比结果。用户创建实验时选择对照组与实验组模型，系统并行执行评测并以可视化图表展示各指标差异，辅助模型迭代决策。

## 6.3 评测指标说明

平台内置多种常用评测指标，覆盖分类、生成与智能体工作流等任务类型。

表 6-1 内置评测指标

| 指标名称 | 适用任务类型 | 取值范围 | 说明 |
| 准确率 | 分类 | 0-1 | 预测正确样本占比 |
| 精确率 | 分类 | 0-1 | 预测为正的样本中正确的比例 |
| 召回率 | 分类 | 0-1 | 正样本被正确召回的比例 |
| F1 值 | 分类 | 0-1 | 精确率与召回率的调和平均 |
| BLEU-4 | 生成 | 0-1 | 译文与参考译文的 n-gram 重合度 |
| ROUGE-L | 摘要 | 0-1 | 基于最长公共子序列的相似度 |
| 任务成功率 | 智能体工作流 | 0-1 | 成功完成任务的比例 |
| 步骤合规率 | 智能体工作流 | 0-1 | 执行步骤符合规范的比例 |
| 工具调用准确率 | 智能体 | 0-1 | 工具选择与参数正确的比例 |
| 平均响应时间 | 性能 | 毫秒 | 单次调用平均耗时 |

## 6.4 报告查看与导出

任务完成后，在任务列表点击"查看报告"。报告包含概览信息、指标明细与样本明细三部分。支持导出为 PDF、Excel、HTML 格式。在"指标报告"页面勾选多个历史任务可对比各指标变化趋势。

# 第七章 插件管理

## 7.1 插件管理

插件管理模块支持以插件方式扩展平台能力。用户可在插件市场浏览可用插件，一键安装或卸载。已安装插件支持启用、停用与版本升级。插件包括自定义指标插件、数据源适配插件、模型适配器插件等类型。

## 7.2 自定义指标开发

用户可按照平台提供的 SDK 规范开发自定义评测指标插件。插件需实现 evaluate 接口，返回标准化的指标结果。开发完成后可通过插件管理页面上传并注册。

# 第八章 计费管理

## 8.1 计费管理

计费管理模块展示平台资源使用情况与费用明细。支持按时间维度查看 API 调用次数、Token 消耗量、存储用量等统计信息。管理员可配置计费规则与预算告警阈值。

## 8.2 用量统计

系统按日、周、月维度自动汇总资源用量，支持导出为 Excel 报表。用量异常时系统自动触发告警通知。

# 第九章 系统设置

## 9.1 系统设置

系统设置页面提供全局参数配置功能，包括：日志保留天数（默认 30 天）、任务最大并发数、模型调用默认超时时间、报告导出水印、邮件通知配置等。修改后点击"保存"，部分参数需重启服务生效。

## 9.2 用户与权限管理

管理员可在"用户管理"页面新增用户、重置密码、停用或启用账号。平台内置三类角色：管理员（全部权限）、评测工程师（数据集、模型、任务、报告操作权限）、访客（只读权限）。权限变更即时生效。

# 第十章 日志管理

## 10.1 运行日志查询

进入"日志管理"页面，可按日志级别（DEBUG、INFO、WARNING、ERROR）、时间范围与关键字组合查询。例如查询关键字"评测配置校验通过"可定位任务配置校验记录。

## 10.2 日志过滤与导出

查询结果支持按级别高亮展示，可点击"导出"将当前查询结果下载为文本文件，便于离线分析与问题上报。

## 10.3 日志归档策略

系统每日定时执行日志归档任务，将超过保留天数（默认 30 天）的日志文件移入归档目录，并记录归档文件数量。归档完成后输出"日志归档完成，共归档 N 个文件"。保留天数可在系统参数中调整。

# 第十一章 常见问题

## 11.1 安装部署类问题

问题：启动服务时提示端口被占用。处理：修改 config.yaml 中的服务端口，或停止占用端口的进程后重启。

问题：数据库初始化失败。处理：检查数据库账号权限与连接地址，确认数据库服务已启动后重新执行 init_db.sh。

## 11.2 使用类问题

问题：上传数据集后校验不通过。处理：查看校验报告，根据错误码 AE-1001 检查字段定义与文件编码是否符合模板要求。

问题：保存评测任务时提示缺少必要字段。处理：确认已完整配置模型名称、数据集路径、指标、输出目录四项内容。

问题：模型连通性测试失败。处理：根据错误码 AE-2001、AE-2002 检查网络、密钥与超时配置。

## 11.3 性能与稳定性问题

问题：大规模任务执行缓慢。处理：在系统参数中提高任务并发数，或按推荐配置扩容服务器。

问题：磁盘空间不足。处理：缩短日志保留天数，或手动执行日志归档清理历史文件。

# 第十二章 附录

## 12.1 错误码表

表 12-1 错误码表

| 错误码 | 说明 | 处理建议 |
| AE-1001 | 数据集格式校验失败 | 检查文件格式与字段定义是否符合模板 |
| AE-1002 | 数据集路径不存在 | 确认存储路径正确且具备读取权限 |
| AE-2001 | 模型连接失败 | 检查网络连通性与 API 密钥配置 |
| AE-2002 | 模型响应超时 | 调整超时参数或检查模型服务状态 |
| AE-3001 | 评测配置缺少必要字段 | 补齐模型名称、数据集路径、指标、输出目录 |
| AE-3002 | 工作流节点配置冲突 | 检查节点依赖关系与参数映射 |
| AE-4001 | 权限不足 | 联系管理员分配相应角色权限 |
| AE-5001 | 日志归档失败 | 检查磁盘空间与归档目录写权限 |

## 12.2 技术支持

如在使用过程中遇到本手册未覆盖的问题，请记录错误码与操作步骤，联系项目技术支持团队获取帮助。
"""

# ========== 排版工具函数 ==========
def set_run_font(run, ascii_name, eastasia_name, size, bold=False):
    run.font.name = ascii_name
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), ascii_name)
    rFonts.set(qn('w:hAnsi'), ascii_name)
    rFonts.set(qn('w:eastAsia'), eastasia_name)
    rPr.append(rFonts)

def set_para_fmt(p, line_pt, space_before=0, space_after=0,
                 first_indent=None, left_indent=None, align=None):
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(line_pt)
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if first_indent is not None:
        fmt.first_line_indent = Pt(first_indent)
    if left_indent is not None:
        fmt.left_indent = Pt(left_indent)
    if align is not None:
        p.alignment = align

def add_h1(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, 30, space_before=12, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(p.add_run(text), 'Times New Roman', '黑体', 16, bold=True)

def add_h2(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, 26, space_before=8, space_after=4)
    set_run_font(p.add_run(text), 'Times New Roman', '黑体', 14, bold=True)

def add_body(doc, text, center=False):
    p = doc.add_paragraph()
    set_para_fmt(p, 22, first_indent=None if center else 24,
                 align=WD_ALIGN_PARAGRAPH.CENTER if center else None)
    set_run_font(p.add_run(text), 'Times New Roman', '宋体', 12)

def add_step(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, 22, left_indent=24)
    set_run_font(p.add_run(text), 'Times New Roman', '宋体', 12)

def add_image(doc, image_path, caption=""):
    if not os.path.exists(image_path):
        print(f"   ⚠️ 图片不存在，跳过: {image_path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_fmt(p, 22, space_before=6, space_after=2)
    run = p.add_run()
    run.add_picture(image_path, width=IMAGE_WIDTH)
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_fmt(cap_p, 18, space_before=2, space_after=6)
        set_run_font(cap_p.add_run(caption), 'Times New Roman', '宋体', 9)

def is_sep_row(row):
    return any('-' in c for c in row) and all(re.fullmatch(r'[-: ]*', c) for c in row)

def add_table(doc, rows):
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = 1
    for r_i, row in enumerate(rows):
        for c_i in range(cols):
            cell_text = row[c_i] if c_i < len(row) else ""
            cell = table.cell(r_i, c_i)
            cell.text = ""
            p = cell.paragraphs[0]
            set_para_fmt(p, 16, space_before=1, space_after=1)
            set_run_font(p.add_run(cell_text), 'Times New Roman', '宋体', 10.5, bold=(r_i == 0))

def add_page_number_footer(doc):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_page = paragraph.add_run()
    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)
    run_page.font.name = 'Times New Roman'
    run_page.font.size = Pt(10)
    run_end = paragraph.add_run(" 页")
    run_end.font.name = 'Times New Roman'
    run_end.font.size = Pt(10)

# ========== 主函数 ==========
def main():
    # 构建截图路径映射
    screenshot_paths = {}
    sd = Path(SCREENSHOT_DIR)
    if sd.exists():
        for section, items in SCREENSHOT_MAP.items():
            screenshot_paths[section] = []
            for fname, caption in items:
                fpath = str(sd / fname)
                screenshot_paths[section].append((fpath, caption))
        found = sum(len(v) for v in screenshot_paths.values())
        print(f"📸 截图目录已加载，共映射 {found} 张截图")
    else:
        print(f"⚠️ 截图目录不存在: {SCREENSHOT_DIR}，将生成无截图版本")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 页眉
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_header = hp.add_run(f"{SOFTWARE_NAME} {VERSION}")
    set_run_font(run_header, 'Times New Roman', '宋体', 12, bold=True)
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    add_page_number_footer(doc)

    # 封面
    p = doc.add_paragraph()
    set_para_fmt(p, 36, space_before=120, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(p.add_run(SOFTWARE_NAME), 'Times New Roman', '黑体', 18, bold=True)
    p = doc.add_paragraph()
    set_para_fmt(p, 30, space_before=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(p.add_run(VERSION), 'Times New Roman', '黑体', 16, bold=True)
    p = doc.add_paragraph()
    set_para_fmt(p, 30, space_before=12, space_after=60, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(p.add_run("用户操作手册"), 'Times New Roman', '黑体', 16, bold=True)
    doc.add_page_break()

    # 解析正文并插入截图
    lines = MANUAL_CONTENT.strip().split('\n')
    i = 0
    current_h2 = ""
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # 表格
        if line.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                if not is_sep_row(cells):
                    rows.append(cells)
                i += 1
            add_table(doc, rows)
            continue

        if line.startswith('# '):
            add_h1(doc, line[2:])
            current_h2 = ""
        elif line.startswith('## '):
            current_h2 = line[3:]
            add_h2(doc, current_h2)
            # 在该节标题后插入对应截图
            if current_h2 in screenshot_paths:
                for fpath, caption in screenshot_paths[current_h2]:
                    add_image(doc, fpath, caption)
        elif re.match(r'^\d+[\.、]', line) or line.startswith('- '):
            add_step(doc, line)
        elif re.match(r'^表\s*\d', line):
            add_body(doc, line, center=True)
        else:
            add_body(doc, line)
        i += 1

    doc.save(OUTPUT_FILE)
    page_est = len([l for l in lines if l.strip()]) // 25 + len(screenshot_paths) * 2
    print(f"\n✅ 说明书生成完成: {OUTPUT_FILE}")
    print(f"   预计约 {page_est} 页（含截图），请打开 Word 确认实际页数 ≥ 15 页")

if __name__ == "__main__":
    main()