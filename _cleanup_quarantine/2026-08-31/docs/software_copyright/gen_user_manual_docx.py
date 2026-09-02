# -*- coding: utf-8 -*-
"""Generate the .docx user manual from the .md source.

Fonts: Chinese -> SimSun (宋体), English -> Courier New.
Adds a page header (页眉) and footer page number (页码).
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"d:\AgentFlow-Eval\docs\software_copyright\AgentFlow-Eval_V1.0.0_User_Manual.md"
OUT = r"d:\AgentFlow-Eval\docs\software_copyright\AgentFlow-Eval_V1.0.0_User_Manual.docx"

FONT_EN = "Courier New"
FONT_CN = "宋体"
SCREENSHOT_DIR = r"C:\Users\yunqi\Desktop\软著截图"


def set_run_font(run, size=11, bold=False):
    run.font.name = FONT_EN
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    rFonts.set(qn("w:eastAsia"), FONT_CN)


def add_rich_text(paragraph, text, size=11):
    """Add text handling **bold** and `code` inline markup."""
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            for cp in re.split(r"(`[^`]*`)", part):
                if not cp:
                    continue
                if cp.startswith("`") and cp.endswith("`"):
                    run = paragraph.add_run(cp[1:-1])
                    set_run_font(run, size=size)
                else:
                    run = paragraph.add_run(cp)
                    set_run_font(run, size=size)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    return run


def add_table(doc, rows):
    ncols = len(rows[0])
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, cell_text in enumerate(row):
            p = cells[ci].paragraphs[0]
            run = p.add_run(cell_text.strip())
            set_run_font(run, size=10, bold=(ri == 0))


def main():
    with open(SRC, "r", encoding="utf-8-sig") as f:
        lines = f.read().splitlines()

    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # header (页眉)
    hp = sec.header.paragraphs[0]
    hp.text = "AgentFlow-Eval 智能体工作流评测平台 V1.0.0"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hp.runs:
        set_run_font(run, size=9)

    # footer page number (页码)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("第 ")
    set_run_font(r1, size=9)
    add_page_number_field(fp)
    r2 = fp.add_run(" 页")
    set_run_font(r2, size=9)

    i = 0
    n = len(lines)
    while i < n:
        stripped = lines[i].strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        # image: ![caption](filename)
        m = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if m:
            caption = m.group(1)
            img_path = os.path.join(SCREENSHOT_DIR, m.group(2))
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_path, width=Cm(13))
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                c_run = cp.add_run(caption)
                set_run_font(c_run, size=9)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("[截图缺失: " + m.group(2) + "]")
                set_run_font(r, size=10)
            i += 1
            continue

        # table
        if stripped.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                add_table(doc, rows)
            continue

        # headings
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            add_rich_text(p, stripped[4:], size=13)
            for r in p.runs:
                r.font.bold = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            add_rich_text(p, stripped[3:], size=15)
            for r in p.runs:
                r.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_rich_text(p, stripped[2:], size=18)
            for r in p.runs:
                r.font.bold = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            add_rich_text(p, stripped.lstrip("> ").strip(), size=10)
            p.paragraph_format.left_indent = Cm(0.5)
            for r in p.runs:
                r.font.italic = True
            i += 1
            continue

        # bullet list
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:], size=11)
            i += 1
            continue

        # numbered list (keep literal numbering to avoid cross-list continuation)
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            add_rich_text(p, stripped, size=11)
            i += 1
            continue

        # normal paragraph
        p = doc.add_paragraph()
        add_rich_text(p, stripped, size=11)
        p.paragraph_format.space_after = Pt(6)
        i += 1

    doc.core_properties.title = "AgentFlow-Eval V1.0.0 用户使用手册"
    doc.core_properties.author = "AgentFlow"

    doc.save(OUT)
    print("DONE:", OUT)


if __name__ == "__main__":
    main()

