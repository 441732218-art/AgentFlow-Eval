# -*- coding: utf-8 -*-
"""Generate the .docx soft-copyright source submission from the frozen .txt.

Reads AgentFlow-Eval_V1.0.0_source_code.txt and emits a .docx with:
- monospace font (Courier New), ~57 lines per page (explicit page breaks)
- header (页眉) + footer page number (页码)
- preserves code order and content (no modification)
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"d:\AgentFlow-Eval\docs\software_copyright\AgentFlow-Eval_V1.0.0_source_code.txt"
OUT = r"d:\AgentFlow-Eval\docs\software_copyright\AgentFlow-Eval_V1.0.0_source_code.docx"
LINES_PER_PAGE = 57
FONT = "Courier New"
EAST_ASIA = "宋体"
FONT_SIZE = 9


def set_font(run, name=FONT, east_asia=EAST_ASIA, size=FONT_SIZE):
    run.font.name = name
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), east_asia)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    return run


def strip_top_block(lines):
    """Remove the leading title line and the first file's '====' header block,
    so the document starts directly with the first file's source code."""
    i = 1  # skip the title line at index 0
    while i < len(lines) and lines[i].strip() == "":
        i += 1
    seen_separator = 0
    while i < len(lines):
        if lines[i].strip() == "================================":
            seen_separator += 1
            i += 1
            if seen_separator == 2:
                break
        else:
            i += 1
    while i < len(lines) and lines[i].strip() == "":
        i += 1
    return lines[i:]


def main():
    with open(SRC, "r", encoding="utf-8-sig") as f:
        lines = strip_top_block(f.read().splitlines())

    doc = Document()

    # page setup: A4
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)

    # header (页眉)
    hp = sec.header.paragraphs[0]
    hp.text = "AgentFlow-Eval 智能体工作流评测平台 V1.0.0"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # footer page number (页码)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("第 ")
    add_page_number_field(fp)
    fp.add_run(" 页")

    # body: one paragraph per line, monospace, page break every 57 lines
    total = len(lines)
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(11)
        run = p.add_run(line)
        set_font(run)
        if (i + 1) % LINES_PER_PAGE == 0 and (i + 1) < total:
            run.add_break(WD_BREAK.PAGE)

    doc.core_properties.title = "AgentFlow-Eval 智能体工作流评测平台 V1.0.0"
    doc.core_properties.author = "AgentFlow"

    doc.save(OUT)
    print("DONE:", OUT)
    print("lines:", total)
    print("pages(约):", (total + LINES_PER_PAGE - 1) // LINES_PER_PAGE)


if __name__ == "__main__":
    main()
