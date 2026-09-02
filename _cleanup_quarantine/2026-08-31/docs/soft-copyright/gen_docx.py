# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pathlib

OUT = pathlib.Path(r"D:\AgentFlow-Eval\copyright_output\软著代码提取教案.docx")
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def h(text, level=1):
    return doc.add_heading(text, level=level)

def p(text, bold=False):
    r = doc.add_paragraph()
    run = r.add_run(text)
    run.bold = bold
    return r

def code(text):
    r = doc.add_paragraph()
    r.paragraph_format.left_indent = Cm(1)
    run = r.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return r

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, x in enumerate(headers):
        t.rows[0].cells[i].text = x
        for pp in t.rows[0].cells[i].paragraphs:
            for rr in pp.runs:
                rr.bold = True
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(v)
    doc.add_paragraph()

doc.add_heading("软著源代码提取实操教案", level=0)
p("小白 30 分钟终端速成指南", bold=True)
doc.add_paragraph()

h("一、教案基本信息")
table(["项目","内容"],[
    ["课程名称","软著源程序鉴别材料 终端一键提取实操"],
    ["适用对象","零基础开发者/行政人员/首次申请软著者"],
    ["预计耗时","25-35 分钟"],
    ["操作系统","Windows 10/11 PowerShell 5.1+"],
    ["前置条件","Python 3.8+ 项目源码在本地"],
    ["最终产出","source_code_60pages.txt + .html"],
])

h("二、教学目标")
p("1. 用5条终端命令完成60页源代码提取")
p("2. 理解每步在做什么")
p("3. 遇到报错能自行排查")
p("4. 产出符合版权中心审查规范的最终文件")

h("三、全流程鸟瞰")
code("Step 0  环境检查 (2 min)")
code("Step 1  创建目录+写入脚本1 (5 min)")
code("Step 2  运行脚本1 生成原始txt (1 min)")
code("Step 3  写入脚本2 (3 min)")
code("Step 4  运行脚本2 裁剪3000有效行 (1 min)")
code("Step 5  写入脚本3 (5 min)")
code("Step 6  运行脚本3 生成60页HTML (1 min)")
code("Step 7  验证+导出PDF (5 min)")

h("四、分步操作")

h("Step 0 环境检查", level=2)
code("python --version")
p("预期: Python 3.8+")
code("echo $PSVersionTable.PSVersion")
p("预期: Major 5 或更高")
table(["现象","解决"],[
    ["提示不是命令","python.org安装 勾选Add to PATH"],
    ["版本<3.8","升级Python"],
])

h("Step 1 创建目录+写入build_source.py", level=2)
code("mkdir D:\\AgentFlow-Eval\\docs\\soft-copyright -Force")
code("mkdir D:\\AgentFlow-Eval\\copyright_output -Force")
p("核心逻辑: 遍历.py 前1500有效行+后1500有效行 拼接输出txt")
code("Test-Path D:\\AgentFlow-Eval\\docs\\soft-copyright\\build_source.py")

h("Step 2 运行脚本1", level=2)
code("python D:\\AgentFlow-Eval\\docs\\soft-copyright\\build_source.py")
p("预期输出:")
code("[INFO] 扫描到 178 个 .py 文件")
code("[OK] 有效行: 3186")
table(["现象","解决"],[
    ["FileNotFoundError","检查ROOT路径"],
    ["有效行<3000","代码不足 全部提交"],
    ["编码报错","加errors=ignore"],
])

h("Step 3 写入trim_fix.py", level=2)
p("作用: 精确裁剪到3000有效行 (60页x50行)")

h("Step 4 运行脚本2", level=2)
code("python D:\\AgentFlow-Eval\\docs\\soft-copyright\\trim_fix.py")
p("预期:")
code("[OK] 裁剪完成: 有效行 3000")
p("必须=3000!", bold=True)

h("Step 5 写入make_html.py", level=2)
p("作用: 3000行按50行/页分页 带页眉页码HTML")
table(["变量","值","说明"],[
    ["SOFTWARE","AgentFlow-Eval","与申请表一致"],
    ["VERSION","V1.0","与申请表一致"],
    ["LINES_PER_PAGE","50","每页有效行"],
])

h("Step 6 运行脚本3", level=2)
code("python D:\\AgentFlow-Eval\\docs\\soft-copyright\\make_html.py")
p("预期:")
code("[INFO] 总页数: 60")
code("[INFO] 每页有效行 min/max = 50/50")
code("[SAVED] -> source_code_60pages.html")

h("Step 7 验证+导出PDF", level=2)
code("start D:\\AgentFlow-Eval\\copyright_output\\source_code_60pages.html")
table(["#","检查项","标准"],[
    ["1","页码","第1页/共60页到第60页/共60页"],
    ["2","页眉","软件全称+版本号每页一致"],
    ["3","连贯性","无乱码无半截函数"],
    ["4","省略标记","第30到31页之间有"],
])
p("导出: Ctrl+P 另存为PDF 纵向 勾选背景图形 保存")

h("五、命令速查卡")
code("python --version")
code("mkdir D:\\AgentFlow-Eval\\docs\\soft-copyright -Force")
code("mkdir D:\\AgentFlow-Eval\\copyright_output -Force")
code("python D:\\AgentFlow-Eval\\docs\\soft-copyright\\build_source.py")
code("python D:\\AgentFlow-Eval\\docs\\soft-copyright\\trim_fix.py")
code("python D:\\AgentFlow-Eval\\docs\\soft-copyright\\make_html.py")
code("start D:\\AgentFlow-Eval\\copyright_output\\source_code_60pages.html")

h("六、换项目只改4个变量")
table(["变量","所在脚本","改成"],[
    ["ROOT","build_source.py","你的项目根目录"],
    ["SOFTWARE","build+make_html","申请表软件全称"],
    ["VERSION","make_html.py","申请表版本号"],
    ["AUTHOR","build_source.py","著作权人"],
])

h("七、FAQ")
table(["问题","解答"],[
    ["不是Python项目","改*.java/*.ts 注释前缀改//"],
    ["代码不足3000行","全部提交 不足60页也合规"],
    ["页眉必须一字不差","是 含空格大小写版本号"],
    ["只交txt行不行","可以 建议交PDF"],
])

h("八、口诀")
pp = doc.add_paragraph()
pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = pp.add_run("三脚本 三回车 六十页 五十行\n页眉名称对申请表 省略标记不能忘\n浏览器里Ctrl+P PDF到手心不慌")
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()
p("教案版本 V1.0 | 2026-07-27")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"[OK] 已生成: {OUT}")
