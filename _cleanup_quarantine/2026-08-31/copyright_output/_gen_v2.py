import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

INPUT = r"D:\AgentFlow-Eval\copyright_output\source_code_60pages_CLEANED.txt"
OUTPUT = r"D:\AgentFlow-Eval\copyright_output\AgentFlow-Eval_源程序鉴别材料_60页_v2.docx"

with open(INPUT, "r", encoding="utf-8") as f:
    raw = [l.rstrip("\r\n") for l in f.readlines()]

def clean_line(line):
    s = line.strip()
    if re.match(r'^\s*#\s*\(?[Cc]\)\s*202[0-9]\s*[\u4e00-\u9fa5]*', s):
        return None
    if s.startswith('"""') and "李凯昕" in s:
        return None
    return line

cleaned = [cl for l in raw if (cl := clean_line(l)) is not None]
print(f"Raw: {len(raw)}  Cleaned: {len(cleaned)}  removed: {len(raw)-len(cleaned)}")

LPP = 56
first = cleaned[:LPP*30]
last = cleaned[-LPP*30:]
sep = ["", "# "+"="*72, "#  软著材料分隔线：以下为源代码后 30 页", "# "+"="*72, ""]
lines = first + sep + last
print(f"Output: {len(lines)} lines  (first={len(first)} + sep={len(sep)} + last={len(last)})")

doc = Document()
for sec in doc.sections:
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(1.8)

s = doc.styles["Normal"]
s.font.name = "Consolas"
s.font.size = Pt(9)
s.font.color.rgb = RGBColor(0, 0, 0)
s.paragraph_format.space_before = Pt(0)
s.paragraph_format.space_after = Pt(0)
s.paragraph_format.line_spacing = Pt(13)
s.paragraph_format.first_line_indent = Cm(0)

# Header
for sec in doc.sections:
    hdr = sec.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    rl = hp.add_run("AgentFlow-Eval V1.0\t\t")
    rl.font.name = "Consolas"
    rl.font.size = Pt(8)
    rl.font.color.rgb = RGBColor(100, 100, 100)
    rp = hp.add_run()
    rp.font.name = "Consolas"
    rp.font.size = Pt(8)
    rp.font.color.rgb = RGBColor(100, 100, 100)
    rp._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    rp._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    rp._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

sep_start = len(first)
for i, line in enumerate(lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(13)
    p.paragraph_format.first_line_indent = Cm(0)
    d = line.replace("\t", "    ") or " "
    run = p.add_run(d)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    if sep_start <= i < sep_start + len(sep):
        run.bold = True
        run.font.color.rgb = RGBColor(80, 80, 80)

print("Saving...")
doc.save(OUTPUT)
print(f"Done! {OUTPUT}")
