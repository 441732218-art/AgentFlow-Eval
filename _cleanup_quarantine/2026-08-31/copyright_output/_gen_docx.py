# -*- coding: utf-8 -*-
"""
软著源代码 Word 生成器 - 合规修正版
严格遵循 CPCC 规范：每页50行 / 页眉页脚 / 前后30页截取
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ==================== 配置参数 ====================
SOURCE_FILE = r"D:\AgentFlow-Eval\copyright_output\source_code_60pages_CLEANED.txt"
OUTPUT_FILE = r"D:\AgentFlow-Eval\copyright_output\AgentFlow-Eval_源程序鉴别材料_60页_20260731.docx"
SOFTWARE_NAME = "AgentFlow-Eval Agent自动化评测工作台 V1.0"
LINES_PER_PAGE = 50          # 硬性要求：每页50行
TOTAL_PAGES = 60             # 总页数
FONT_NAME = "Consolas"       # 等宽字体
FONT_SIZE = Pt(9)            # 小五号
LINE_SPACING = Pt(13)        # 固定行距13pt
MARGIN_TOP = Cm(2.0)
MARGIN_BOTTOM = Cm(1.8)
MARGIN_LEFT = Cm(2.2)
MARGIN_RIGHT = Cm(1.8)

# ==================== 代码清洗规则 ====================
def clean_line(line):
    """过滤冗余版权注释、作者声明等非功能性代码行"""
    stripped = line.strip()
    
    # 删除 (c) 2026 李凯昕 及类似版权行
    if re.match(r'^\s*#\s*[\(（]?[Cc][\)）]?\s*202[0-9]\s*[\u4e00-\u9fa5]*', stripped):
        return None
    
    # 删除纯作者声明行
    if re.match(r'^\s*#\s*[Aa]uthor\s*:', stripped):
        return None
    
    # 删除文件头部的三引号文档字符串中的版权信息（简化处理）
    if '李凯昕' in stripped and ('"""' in stripped or "'''" in stripped):
        return None
        
    return line


# ==================== 主逻辑 ====================
def main():
    # 1. 读取并清洗源码
    print(f"Reading source: {SOURCE_FILE}")
    with open(SOURCE_FILE, 'r', encoding='utf-8') as f:
        raw_lines = f.readlines()
    
    cleaned_lines = []
    for line in raw_lines:
        result = clean_line(line)
        if result is not None:
            cleaned_lines.append(result.rstrip('\n'))
    
    print(f"Raw lines: {len(raw_lines)}, Cleaned lines: {len(cleaned_lines)}")
    
    # 2. 精确截取前后30页
    # 安全截取：从1500行向前搜索顶层定义边界
    first_count = LINES_PER_PAGE * (TOTAL_PAGES // 2)
    for j in range(first_count, max(0, first_count - 200), -1):
        stripped = cleaned_lines[j].strip()
        if stripped.startswith(("def ", "class ")) or (stripped and not cleaned_lines[j][0].isspace()):
            first_count = j
            print(f"  ✅ 安全分割点调整至第{j}行")
            break
    last_count = LINES_PER_PAGE * (TOTAL_PAGES // 2)
    
    if len(cleaned_lines) >= first_count + last_count:
        first_part = cleaned_lines[:first_count]
        last_part = cleaned_lines[-last_count:]
        print(f"截取模式: 前{first_count}行 + 后{last_count}行")
    else:
        # 不足60页时全量保留
        first_part = cleaned_lines
        last_part = []
        print(f"警告: 源码仅{len(cleaned_lines)}行，不足60页，将全量输出")
    
    # 3. 构建分隔线（仅3行）
    separator = [
        "# " + "━" * 52,
        "# 【软著材料分隔线】以下为源代码后30页（共1500行）",
        "# " + "━" * 52,
    ]
    
    final_lines = first_part + separator + last_part
    print(f"最终写入行数: {len(final_lines)}")
    
    # 4. 创建 Word 文档
    doc = Document()
    
    # 设置页面布局
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    
    # 设置段落格式（固定行距 + 无缩进）
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0)
    
    # 5. 添加合规页眉
    header = section.header
    header.is_linked_to_previous = False
    hdr_para = header.paragraphs[0]
    hdr_para.text = SOFTWARE_NAME
    hdr_run = hdr_para.runs[0]
    hdr_run.font.name = FONT_NAME
    hdr_run.font.size = Pt(9)
    hdr_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 6. 添加合规页脚（自动页码域）
    footer = section.footer
    footer.is_linked_to_previous = False
    ftr_para = footer.paragraphs[0]
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run_pre = ftr_para.add_run("第 ")
    run_pre.font.name = FONT_NAME
    run_pre.font.size = Pt(9)
    
    # 插入 PAGE 域代码
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    fld_char_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    
    run_field = ftr_para.add_run()
    run_field._element.append(fld_char_begin)
    run_field2 = ftr_para.add_run()
    run_field2._element.append(instr_text)
    run_field3 = ftr_para.add_run()
    run_field3._element.append(fld_char_separate)
    run_page = ftr_para.add_run("1")  # 占位符，Word打开后自动更新
    run_page.font.name = FONT_NAME
    run_page.font.size = Pt(9)
    run_field4 = ftr_para.add_run()
    run_field4._element.append(fld_char_end)
    
    run_suf = ftr_para.add_run(" 页")
    run_suf.font.name = FONT_NAME
    run_suf.font.size = Pt(9)
    
    # 7. 写入代码内容
    total = len(final_lines)
    for i, line in enumerate(final_lines):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        
        # 分隔线加粗灰色标识
        if i >= len(first_part) and i < len(first_part) + len(separator):
            run.bold = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        if (i + 1) % 500 == 0:
            print(f"  ... {i+1}/{total}")
    
    # 8. 保存文档
    doc.save(OUTPUT_FILE)
    file_size = os.path.getsize(OUTPUT_FILE)
    print(f"\nDone! Saved: {OUTPUT_FILE}")
    print(f"   Size: {file_size:,} bytes")
    print(f"   Lines: {total}, Target pages: {TOTAL_PAGES} ({LINES_PER_PAGE} lines/page)")


if __name__ == "__main__":
    main()



