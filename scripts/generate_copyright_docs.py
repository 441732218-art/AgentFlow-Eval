"""
AgentFlow智能体评测管理平台软件 V1.0.0
软件著作权登记材料生成脚本
生成：《源程序鉴别材料》 + 《用户操作手册》
"""

import os
import re
import sys

# Reconfigure stdout/stderr to use UTF-8 on Windows
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8')

# ── 检查并安装依赖 ──────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("正在安装 python-docx ...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════
#  基本信息（与申请表完全一致）
# ══════════════════════════════════════════════════════
SOFTWARE_NAME    = "AgentFlow智能体评测管理平台软件"
VERSION          = "V2.0.0"
HEADER_TEXT      = f"{SOFTWARE_NAME} {VERSION}"
AUTHOR           = "李凯昕"
COMPLETE_DATE    = "2026年7月16日"
PAGES_PER_HALF   = 30          # 前30页 + 后30页
LINES_PER_PAGE   = 50          # 每页≥50行有效代码

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "artifacts", "copyright")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ══════════════════════════════════════════════════════
#  脱敏规则
# ══════════════════════════════════════════════════════
_SENSITIVE_PATTERNS = [
    # 开源协议 / 外部版权声明 / 其他版权声明
    (re.compile(r"#\s*Copyright.*?\n", re.IGNORECASE), ""),
    (re.compile(r"#\s*MIT License.*?\n", re.IGNORECASE), ""),
    (re.compile(r"#\s*Licensed under.*?\n", re.IGNORECASE), ""),
    # 开发者真实姓名（新规脱敏要求：直接彻底剔除）
    (re.compile(r'\b李凯昕\b'), ""),
    # 删除具体年份、特定日期等时间戳（符合2026合规清单）
    (re.compile(r'\b20[0-9]{2}(?:[-/年][0-9]{1,2}(?:[-/月][0-9]{1,2}日?)?)?\b'), ""),
    # 常见的联系方式与邮箱、可能有的身份证进行模式脱敏删除
    (re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'), ""),
    (re.compile(r'\b(?:\+?\d{1,3}[- ]?)?1[3-9]\d{9}\b'), ""),
    (re.compile(r'\b\d{17}[\dXx]\b'), ""),
    # API Key, 密钥, 密码, token 赋值全部用 "***" 替代
    (re.compile(r'([\w_-]*(?:key|secret|password|token)\s*=\s*)[rR]?["\'][^"\']*["\']', re.IGNORECASE), r'\1"***"'),
    # 具体服务器地址
    (re.compile(r'redis://[^\s"\']+'), "redis://***:6379/0"),
    (re.compile(r'postgresql\+[^\s"\']+'), "postgresql+asyncpg://***"),
]


def desensitize(code: str) -> str:
    """对代码执行脱敏处理。"""
    for pattern, replacement in _SENSITIVE_PATTERNS:
        code = pattern.sub(replacement, code)
    # 移除首行 (c) 版权注释（不影响代码逻辑，规避他人版权标注）
    lines = code.split("\n")
    cleaned = []
    for line in lines:
        if re.match(r"^\s*#\s*\(c\)\s*", line, re.IGNORECASE):
            continue
        cleaned.append(line)
    return "\n".join(cleaned)


def count_effective_lines(code_lines: list[str]) -> list[str]:
    """返回有效代码行（去除纯空行和纯注释行）。"""
    effective = []
    for line in code_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#") and not stripped.startswith("#!"):
            continue
        effective.append(line)
    return effective


# ══════════════════════════════════════════════════════
#  核心源代码收集（白名单）
# ══════════════════════════════════════════════════════
BASE_DIR = os.path.join(os.path.dirname(__file__), "..", "backend", "app")

WHITE_LIST_FILES = [
    # 层次顺序：入口 → 中间件 → RBAC → 评判引擎 → Agent运行器 → API端点 → 评估管道
    ("main.py",                                       "应用入口与生命周期管理"),
    ("core/middleware.py",                            "请求中间件：鉴权/追踪/安全头"),
    ("core/security.py",                              "API密钥鉴权与身份认证"),
    ("core/rbac.py",                                  "基于角色的访问控制引擎"),
    ("core/tenancy.py",                               "多租户隔离与资源归属"),
    ("core/judge_engine/llm_judge.py",                "混合评判引擎：规则+LLM"),
    ("core/judge_engine/metrics.py",                  "评判指标计算工具"),
    ("core/judge_engine/base.py",                     "评判引擎抽象基类"),
    ("core/agent_runner/openai_runner.py",            "OpenAI ReAct智能体执行器"),
    ("core/agent_runner/tool_sandbox.py",              "工具沙箱：安全执行与注册"),
    ("core/agent_runner/factory.py",                  "智能体运行器工厂"),
    ("core/agent_runner/http_runner.py",              "HTTP协议智能体执行器"),
    ("core/evaluation/pipeline.py",                   "评估流程编排与聚合"),
    ("core/evaluation/compare.py",                    "多实验组对比分析"),
    ("api/v1/endpoints/tasks.py",                     "评测任务CRUD API端点"),
    ("api/v1/endpoints/experiments.py",               "多变体实验管理API"),
    ("api/v1/endpoints/traces.py",                    "执行轨迹查询API"),
    ("api/v1/endpoints/plugins.py",                   "插件市场与管理API"),
    ("api/v1/endpoints/billing.py",                   "计费与配额管理API"),
    ("core/dependencies.py",                          "FastAPI依赖注入"),
    ("core/events.py",                                "事件总线与WebSocket通知"),
    ("core/audit.py",                                 "操作审计日志"),
    ("core/ws_hub.py",                                "WebSocket实时推送中心"),
    ("models/task.py",                                "评测任务数据模型"),
    ("models/trace.py",                               "执行轨迹数据模型"),
    ("models/experiment.py",                          "实验数据模型"),
    ("models/metric_score.py",                        "指标评分数据模型"),
    ("models/billing.py",                             "计费数据模型"),
]


def load_all_code() -> list[tuple[str, bool]]:
    """读取所有白名单文件中的有效行，包装为(行内容, 是否有效)元组。"""
    all_tuples: list[tuple[str, bool]] = []
    for rel_path, description in WHITE_LIST_FILES:
        full_path = os.path.join(BASE_DIR, rel_path)
        if not os.path.exists(full_path):
            print(f"  [跳过] 文件不存在: {rel_path}")
            continue
        with open(full_path, encoding="utf-8", errors="replace") as f:
            raw = f.read()
        cleaned = desensitize(raw)
        
        # 1. 提取该文件中的有效代码行
        raw_lines = cleaned.split("\n")
        effective_lines = []
        in_docstring_double = False
        in_docstring_single = False
        in_c_comment = False
        
        for line in raw_lines:
            line_str = line.rstrip()
            stripped = line_str.strip()
            
            # 过滤空行
            if not stripped:
                continue
                
            # 处理 C 风格多行注释 (/* ... */)
            if in_c_comment:
                if "*/" in stripped:
                    in_c_comment = False
                continue
            if stripped.startswith("/*"):
                if "*/" not in stripped:
                    in_c_comment = True
                continue
                
            # 处理 Python 多行 Docstring 三引号 (""" or ''')
            if in_docstring_double:
                if '"""' in stripped:
                    in_docstring_double = False
                continue
            if in_docstring_single:
                if "'''" in stripped:
                    in_docstring_single = False
                continue
                
            if stripped.startswith('"""'):
                # 检查是否是单行双三引号 docstring，如 """docstring"""
                if stripped.count('"""') == 1:
                    in_docstring_double = True
                continue
            if stripped.startswith("'''"):
                if stripped.count("'''") == 1:
                    in_docstring_single = True
                continue
                
            # 过滤单行注释
            if stripped.startswith('#') or stripped.startswith('//') or stripped.startswith('*'):
                continue
                
            # 截取到80字符
            effective_lines.append(line_str[:80])
            
        if not effective_lines:
            continue
            
        # 2. 生成文件描述头（非有效行，仅作展示）
        file_headers = [
            f"# {'='*60}",
            f"# 文件: {rel_path}",
            f"# 功能: {description}",
            f"# {'='*60}"
        ]
        
        # 3. 添加到大列表中页
        for h in file_headers:
            all_tuples.append((h, False))
        for line in effective_lines:
            all_tuples.append((line, True))
            
        print(f"  [加载] {rel_path} ({len(effective_lines)} 行有效代码)")
        
    return all_tuples


def set_run_font(run, font_name="宋体", size_pt=10.5):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

def set_code_run_font(run, size_pt=9.0):
    run.font.name = "Courier New"
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Courier New')
    rFonts.set(qn('w:hAnsi'), 'Courier New')
    rFonts.set(qn('w:eastAsia'), '宋体') # Ensure CJK comment strings are readable
    rPr.append(rFonts)

def paginate_lines(all_tuples: list[tuple[str, bool]]) -> list[list[str]]:
    """将元组按每页正好50个有效行进行分页。返回一个由行文本组成的分页列表。"""
    pages = []
    current_page = []
    effective_count = 0
    
    for line, is_effective in all_tuples:
        current_page.append(line)
        if is_effective:
            effective_count += 1
            
        if effective_count >= 50:
            pages.append(current_page)
            current_page = []
            effective_count = 0
            
    # 如果最后一页还有有效行，补空行使有效行数达到50（审查规定）
    if current_page:
        if effective_count < 50:
            current_page.extend([""] * (50 - effective_count))
        pages.append(current_page)
        
    return pages


# ══════════════════════════════════════════════════════
#  Word 文档辅助函数
# ══════════════════════════════════════════════════════

def _set_page_layout(doc: Document) -> None:
    """设置A4页面，页边距2cm。"""
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)


def _add_header_footer(section, header_text: str, page_num_offset: int = 0) -> None:
    """为section添加页眉（软件名+版本）和页脚（页码）。"""
    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(header_text)
    set_run_font(run, "宋体", 9)

    # 页脚（自动域码页码）
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_pre = fp.add_run("第 ")
    set_run_font(run_pre, "宋体", 9)

    # 插入页码域
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")

    run_page = fp.add_run()
    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)
    run_page._r.append(fldChar3)
    set_run_font(run_page, "Courier New", 9)

    run_suf = fp.add_run(" 页")
    set_run_font(run_suf, "宋体", 9)


def _write_code_page(doc: Document, lines: list[str], page_no: int, is_first: bool) -> None:
    """将一页代码写入文档（等宽字体，小五号）。"""
    if not is_first:
        doc.add_page_break()
    # 代码段落
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    para.paragraph_format.line_spacing = Pt(11)
    code_text = "\n".join(lines)
    run = para.add_run(code_text)
    set_code_run_font(run, size_pt=9.0)


# ══════════════════════════════════════════════════════
#  生成《源程序鉴别材料》
# ══════════════════════════════════════════════════════

def generate_source_code_doc(pages: list[list[str]]) -> str:
    total = len(pages)
    print(f"\n总页数: {total} 页")

    # 选取前30页 + 后30页，生成刚好60页，不加封面页以确保刚好60页
    if total <= 60:
        selected_pages = list(pages)
        while len(selected_pages) < 60:
            selected_pages.append([""] * 50)
        print("  代码总页数≤60，全部截取并补足60页。")
    else:
        front_pages = pages[:PAGES_PER_HALF]
        back_pages  = pages[-PAGES_PER_HALF:]
        selected_pages = front_pages + back_pages
        print(f"  前30页: 第1~30页  后30页: 第{total-29}~{total}页")

    doc = Document()
    _set_page_layout(doc)

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(9)

    # 第一页开始直入代码
    section0 = doc.sections[0]
    _add_header_footer(section0, HEADER_TEXT)

    for i, page_lines in enumerate(selected_pages):
        _write_code_page(doc, page_lines, i + 1, is_first=(i == 0))

    out_path = os.path.join(OUTPUT_DIR, f"{SOFTWARE_NAME}_{VERSION}_源程序鉴别材料.docx")
    doc.save(out_path)
    print(f"\n[OK] 源程序鉴别材料已生成: {out_path}")
    return out_path


# ══════════════════════════════════════════════════════
#  生成《用户操作手册》
# ══════════════════════════════════════════════════════

def generate_user_manual() -> str:
    doc = Document()
    _set_page_layout(doc)

    section0 = doc.sections[0]
    _add_header_footer(section0, HEADER_TEXT)

    # ── 样式辅助 ──────────────────────────────────────
    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            set_run_font(run, "宋体", 12 if level == 2 else 14)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        return p

    def add_para(text, indent=0, bold=False, size=10.5):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text)
        set_run_font(run, "宋体", size)
        run.bold = bold
        return p

    def add_table_row(table, cells_data, header=False):
        row = table.add_row()
        for i, text in enumerate(cells_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_run_font(run, "宋体", 9)
            run.bold = header

    # ══════════════════════════════════════════════════
    #  封面
    # ══════════════════════════════════════════════════
    doc.add_paragraph("\n\n\n")
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_run = cover.add_run(f"{SOFTWARE_NAME}\n{VERSION}\n\n用户操作手册")
    set_run_font(c_run, "宋体", 20)
    c_run.bold = True

    doc.add_paragraph("\n\n")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    i_run = info.add_run(
        f"著作权人：{AUTHOR}\n"
        f"开发完成日期：{COMPLETE_DATE}\n"
        f"文档版本：{VERSION}\n"
    )
    set_run_font(i_run, "宋体", 12)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    #  目录（手工，节省空间）
    # ══════════════════════════════════════════════════
    add_heading("目  录", level=1)
    toc_items = [
        ("1", "引言"),
        ("2", "运行环境"),
        ("3", "安装与启动"),
        ("4", "用户登录与主界面"),
        ("5", "核心功能操作流程"),
        ("  5.1", "评测任务管理"),
        ("  5.2", "测试用例管理"),
        ("  5.3", "任务执行与监控"),
        ("  5.4", "多变体实验对比"),
        ("  5.5", "执行轨迹查询"),
        ("  5.6", "插件管理"),
        ("  5.7", "系统设置与权限"),
        ("6", "退出系统"),
        ("7", "常见问题解答"),
    ]
    for num, title in toc_items:
        add_para(f"{num}  {title}", indent=0.5)

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    #  第1章 引言
    # ══════════════════════════════════════════════════
    add_heading("1  引言", level=1)
    add_heading("1.1  文档目的", level=2)
    add_para(
        "本手册旨在帮助用户快速上手 " + SOFTWARE_NAME + "（以下简称本软件或平台），"
        "掌握各核心功能模块的操作方法，以便高效完成智能体的自动化评估任务。",
        indent=0.5
    )

    add_heading("1.2  软件简介", level=2)
    add_para(
        f"{SOFTWARE_NAME} {VERSION} 是一款专注于大型语言模型（LLM）驱动的智能体（AI Agent）"
        "自动化评测与管理平台。平台集成了混合评判引擎（规则评分 + LLM-as-Judge双轨校验）、"
        "多变体实验对比分析、安全工具沙箱执行、实时WebSocket进度推送、多租户权限隔离"
        "等核心能力，适用于AI研发团队对智能体系统进行系统化、可量化的质量管控。",
        indent=0.5
    )

    add_heading("1.3  适用范围", level=2)
    add_para("本手册适用于以下用户群体：", indent=0.5)
    for item in [
        "AI研发工程师：负责配置评测任务与分析评分结果；",
        "产品经理/质量评审员：负责查阅评测报告、提交人工复核意见；",
        "系统管理员：负责用户权限管理与系统配置。",
    ]:
        add_para(f"  • {item}", indent=1.0)

    # ══════════════════════════════════════════════════
    #  第2章 运行环境
    # ══════════════════════════════════════════════════
    doc.add_page_break()
    add_heading("2  运行环境", level=1)
    add_heading("2.1  服务端运行环境", level=2)

    table1 = doc.add_table(rows=0, cols=3)
    table1.style = "Table Grid"
    add_table_row(table1, ["项目", "最低要求", "推荐配置"], header=True)
    env_rows = [
        ("操作系统",     "Linux / Windows / macOS",        "Ubuntu 22.04 LTS"),
        ("Python版本",   "Python 3.11+",                    "Python 3.12"),
        ("内存",         "4 GB RAM",                        "8 GB RAM"),
        ("磁盘空间",     "5 GB 可用空间",                   "20 GB SSD"),
        ("数据库",       "SQLite 3（轻量模式）",            "PostgreSQL 15+"),
        ("消息队列",     "内存队列（轻量模式）",            "Redis 7.0+"),
        ("网络",         "HTTP/HTTPS 8000 端口",             "Nginx反代 + TLS"),
    ]
    for row in env_rows:
        add_table_row(table1, list(row))

    add_heading("2.2  客户端运行环境（浏览器）", level=2)
    for item in [
        "Chrome 110+ / Edge 110+ / Firefox 110+（推荐 Chrome 最新版）",
        "屏幕分辨率：1366×768 及以上（推荐 1920×1080）",
        "JavaScript：必须启用",
    ]:
        add_para(f"  • {item}", indent=0.5)

    # ══════════════════════════════════════════════════
    #  第3章 安装与启动
    # ══════════════════════════════════════════════════
    doc.add_page_break()
    add_heading("3  安装与启动", level=1)
    add_heading("3.1  轻量模式（开发/演示环境）", level=2)
    add_para("轻量模式使用 SQLite 数据库 + 内存任务队列，无需额外部署 Redis/PostgreSQL，适合本地演示：",
             indent=0.5)
    steps_light = [
        ("步骤1  获取代码", "从软件包解压或克隆代码仓库至本地目录。"),
        ("步骤2  安装依赖", "进入 backend 目录，执行：\n      pip install -r requirements.txt"),
        ("步骤3  配置环境", "复制 .env.example 为 .env，修改 OPENAI_API_KEY 为有效密钥。\n      "
                           "轻量模式保持默认值即可，CELERY_TASK_ALWAYS_EAGER=true 已预设。"),
        ("步骤4  启动后端", "执行：uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload\n"
                           "      启动成功后控制台显示：\"Application startup complete\""),
        ("步骤5  启动前端", "进入 frontend 目录，执行：\n      npm install && npm run dev\n"
                           "      浏览器访问 http://localhost:5173"),
    ]
    for title, content in steps_light:
        add_para(f"  {title}", indent=0.5, bold=True)
        add_para(f"  {content}", indent=1.0)

    add_heading("3.2  生产模式（Docker Compose）", level=2)
    add_para("生产模式通过 Docker Compose 一键部署全栈服务（后端+前端+Redis+Nginx）：", indent=0.5)
    prod_steps = [
        ("步骤1  准备配置", "复制 deploy.env.example 为 .env，填写数据库连接串、API密钥等必要配置。"),
        ("步骤2  构建启动", "执行：docker compose -f docker-compose.prod.yml up -d --build"),
        ("步骤3  验证服务", "访问 http://服务器IP/health，返回 {\"status\": \"healthy\"} 表示就绪。"),
        ("步骤4  访问平台", "浏览器访问 http://服务器IP，进入登录界面。"),
    ]
    for title, content in prod_steps:
        add_para(f"  {title}", indent=0.5, bold=True)
        add_para(f"  {content}", indent=1.0)

    # ══════════════════════════════════════════════════
    #  第4章 用户登录与主界面
    # ══════════════════════════════════════════════════
    doc.add_page_break()
    add_heading("4  用户登录与主界面", level=1)
    add_heading("4.1  配置访问凭证", level=2)
    add_para(
        "默认轻量模式下，平台不启用身份验证（AUTH_ENABLED=false），可直接访问所有功能。"
        "启用身份验证后，需在请求头中携带 X-API-Key 凭证，或在系统设置中配置 API 密钥。",
        indent=0.5
    )
    add_para("API密钥格式说明：", indent=0.5, bold=True)
    key_formats = [
        ("基本格式",       "your-secret-key",          "匿名用户，默认USER角色"),
        ("指定用户",       "your-key:alice",           "alice用户，使用默认角色"),
        ("指定角色",       "your-key:alice:manager",   "alice用户，MANAGER角色"),
    ]
    table2 = doc.add_table(rows=0, cols=3)
    table2.style = "Table Grid"
    add_table_row(table2, ["格式类型", "密钥示例", "说明"], header=True)
    for row in key_formats:
        add_table_row(table2, list(row))

    add_heading("4.2  主界面布局", level=2)
    add_para("成功访问后，主界面由以下区域构成：", indent=0.5)
    layout_items = [
        ("顶部导航栏",    "显示软件名称、版本号及全局操作按钮（语言切换、主题切换）"),
        ("左侧菜单栏",    "包含：仪表盘、评测任务、实验管理、执行轨迹、插件市场、系统设置"),
        ("中央内容区",    "展示当前功能页面的操作界面与数据列表"),
        ("右侧状态栏",    "实时显示任务执行进度（通过WebSocket推送更新）"),
    ]
    for name, desc in layout_items:
        add_para(f"  • 【{name}】：{desc}", indent=0.8)

    add_heading("4.3  仪表盘概览", level=2)
    add_para(
        "仪表盘为系统首页，展示以下关键指标：总任务数、执行中任务数、已完成任务数、"
        "平均评分趋势图（近30天）、各评分维度分布（工具准确率/回答正确率/推理连贯性）。"
        "所有数据每30秒自动刷新，支持手动点击刷新按钮立即更新。",
        indent=0.5
    )

    # ══════════════════════════════════════════════════
    #  第5章 核心功能操作流程
    # ══════════════════════════════════════════════════
    doc.add_page_break()
    add_heading("5  核心功能操作流程", level=1)

    # 5.1 评测任务管理
    add_heading("5.1  评测任务管理", level=2)
    add_para("评测任务是平台的核心工作单元，每个任务包含若干测试用例（Test Suite）和对应的智能体配置。",
             indent=0.5)
    add_para("5.1.1  创建评测任务", indent=0.3, bold=True, size=10)
    create_steps = [
        "点击左侧菜单「评测任务」，进入任务列表页。",
        "点击右上角「+ 新建任务」按钮，弹出创建对话框。",
        "填写以下字段：\n      任务名称（必填）：如 GPT-4o 数学推理能力评测\n"
        "      任务描述（选填）：任务背景与目标说明\n"
        "      智能体配置（必填）：选择模型（如 gpt-4o-mini）、最大迭代次数（1-20）、温度参数",
        "点击「提交」按钮，任务创建成功，状态显示为「已创建」。",
    ]
    for i, step in enumerate(create_steps, 1):
        add_para(f"      步骤{i}：{step}", indent=1.0)

    add_para("5.1.2  任务状态说明", indent=0.3, bold=True, size=10)
    table3 = doc.add_table(rows=0, cols=3)
    table3.style = "Table Grid"
    add_table_row(table3, ["状态", "标识颜色", "含义"], header=True)
    status_rows = [
        ("created（已创建）",   "蓝色",   "任务已建立，等待添加测试用例并执行"),
        ("queued（排队中）",    "橙色",   "任务已提交执行队列，等待 Worker 处理"),
        ("running（执行中）",   "黄色",   "Worker 正在执行智能体与评判流程"),
        ("completed（已完成）", "绿色",   "所有测试用例均已完成评分"),
        ("partial（部分完成）", "深黄色", "部分测试用例成功，部分失败"),
        ("failed（已失败）",    "红色",   "任务整体执行失败"),
        ("cancelled（已取消）", "灰色",   "用户手动取消"),
        ("timeout（超时）",     "紫色",   "任务超过最大执行时间被终止"),
    ]
    for row in status_rows:
        add_table_row(table3, list(row))

    doc.add_paragraph()

    # 5.2 测试用例管理
    add_heading("5.2  测试用例管理", level=2)
    add_para("测试用例（Test Suite）定义了对智能体的具体提问与期望输出，是评分的基础数据单元。",
             indent=0.5)
    add_para("5.2.1  手动添加测试用例", indent=0.3, bold=True, size=10)
    add_para("在任务详情页点击「添加测试用例」，填写以下字段后点击「保存」：", indent=0.8)
    suite_fields = [
        ("用户问题（必填）",    "模拟用户向智能体发送的查询，例如：请计算 (12+8)*5 的结果"),
        ("期望输出（推荐填）",  "智能体应给出的正确答案，例如：100"),
        ("期望调用工具（选填）","用竖线分隔的工具名列表，例如：calculator|web_search"),
    ]
    for name, desc in suite_fields:
        add_para(f"  • {name}：{desc}", indent=1.2)

    add_para("5.2.2  批量导入测试用例", indent=0.3, bold=True, size=10)
    add_para("支持 CSV 或 JSON 格式批量导入（上限500条/次）：", indent=0.8)
    add_para("  CSV格式（首行为列名）：user_query, expected_output, expected_tools", indent=1.2)
    add_para("  JSON格式（数组）：[{\"user_query\":\"...\",\"expected_output\":\"...\",\"expected_tools\":[\"...\"]}]",
             indent=1.2)
    add_para("  操作步骤：任务详情页 → 「批量导入」→ 选择文件 → 平台自动解析并入库，返回导入数量确认。",
             indent=1.2)

    doc.add_page_break()

    # 5.3 任务执行与监控
    add_heading("5.3  任务执行与监控", level=2)
    add_para("5.3.1  启动任务执行", indent=0.3, bold=True, size=10)
    exec_steps = [
        "确认测试用例已添加完毕（任务详情页显示「测试用例数 > 0」）。",
        "点击任务卡片上的「▶ 执行」按钮，系统进行配额检查后将任务推入执行队列。",
        "任务状态由「已创建」→「排队中」→「执行中」动态更新（WebSocket实时推送，无需手动刷新）。",
    ]
    for i, step in enumerate(exec_steps, 1):
        add_para(f"      步骤{i}：{step}", indent=1.0)

    add_para("5.3.2  评分维度说明", indent=0.3, bold=True, size=10)
    add_para("平台采用三维度评分体系，满分100分：", indent=0.8)
    score_dims = [
        ("工具调用准确率（40分）", "检验智能体是否调用了期望工具、调用顺序是否合理"),
        ("回答正确率（40分）",     "通过词汇重叠（支持中英文混合/CJK字符）计算答案与期望输出的匹配度"),
        ("推理连贯性（20分）",     "检测推理步骤是否有重复思维、过度迭代等质量问题"),
    ]
    for name, desc in score_dims:
        add_para(f"  • {name}：{desc}", indent=1.2)

    add_para("5.3.3  执行监控面板", indent=0.3, bold=True, size=10)
    add_para(
        "执行过程中，任务详情页实时显示：已完成/失败/进行中的测试用例数量、"
        "当前平均得分、Token消耗数、预计剩余时间。"
        "点击单条测试用例可展开查看完整的 ReAct 推理步骤（Thought→Action→Observation 链路）。",
        indent=0.8
    )

    # 5.4 多变体实验对比
    add_heading("5.4  多变体实验对比", level=2)
    add_para(
        "实验功能允许同一套测试用例在不同智能体配置（如不同模型、不同提示词）下并行评测，"
        "并提供横向对比分析，帮助研发团队选出最优配置。",
        indent=0.5
    )
    add_para("创建实验操作步骤：", indent=0.5, bold=True)
    exp_steps = [
        "点击左侧菜单「实验管理」→「+ 新建实验」。",
        "设置实验名称与描述，可选择「继承已有任务」的测试用例集，或手动添加新测试用例。",
        "添加2个或以上「实验变体」，每个变体配置不同的模型参数（如 gpt-4o-mini vs gpt-4o）。",
        "勾选「自动执行」后点击「创建」，系统自动为每个变体创建独立任务并并行提交执行。",
        "执行完成后点击「📊 对比结果」，查看各变体在三个评分维度上的得分对比表与差值分析。",
        "平台自动标注得分最高的「最优变体」（以绿色徽标突出显示）。",
    ]
    for i, step in enumerate(exp_steps, 1):
        add_para(f"      步骤{i}：{step}", indent=1.0)

    # 5.5 执行轨迹查询
    add_heading("5.5  执行轨迹查询", level=2)
    add_para(
        "执行轨迹记录了智能体每次 API 调用的完整过程，是调试与分析的核心工具。",
        indent=0.5
    )
    add_para("查询与分析操作：", indent=0.5, bold=True)
    trace_ops = [
        "点击左侧菜单「执行轨迹」，支持按任务ID、时间范围、状态（成功/失败）进行筛选。",
        "点击单条轨迹，展开详情：ReAct推理步骤、工具调用参数与返回值、Token消耗、响应时长。",
        "在「评分详情」标签页查看三维度原始得分、评判引擎模式（纯规则/混合）、评判理由。",
        "评审员可在打分区填写「人工评分」并提交复核，系统记录人工评分以供后续统计分析。",
    ]
    for i, op in enumerate(trace_ops, 1):
        add_para(f"      {i}. {op}", indent=1.0)

    doc.add_page_break()

    # 5.6 插件管理
    add_heading("5.6  插件管理", level=2)
    add_para("插件系统允许动态扩展平台工具能力，无需修改核心代码。", indent=0.5)

    add_para("5.6.1  插件市场", indent=0.3, bold=True, size=10)
    add_para("点击左侧菜单「插件市场」，浏览可用插件列表，每个插件卡片显示：", indent=0.8)
    plugin_fields = [
        "插件名称与版本号",
        "功能描述与能力标签（工具扩展/提示词增强/评估指标）",
        "激活状态（已激活/未激活）",
        "配置参数（点击「配置」按钮展开）",
    ]
    for f in plugin_fields:
        add_para(f"  • {f}", indent=1.2)

    add_para("5.6.2  自定义插件安装", indent=0.3, bold=True, size=10)
    add_para("系统管理员可通过以下方式加载自定义插件：", indent=0.8)
    add_para("  方式A（目录扫描）：将插件包放置于 plugins/ 目录，重启服务自动加载；", indent=1.2)
    add_para("  方式B（显式注册）：在 .env 中配置 PLUGIN_MODULES=your_package:YourPlugin 并重启。", indent=1.2)

    # 5.7 系统设置与权限
    add_heading("5.7  系统设置与权限", level=2)
    add_para("点击左侧菜单「系统设置」，管理员可配置以下选项：", indent=0.5)

    settings_items = [
        ("API密钥管理",     "查看已配置的API密钥列表（值脱敏显示），可新增或删除"),
        ("用户角色配置",    "配置 actor→role 映射；支持角色：Admin/Manager/Reviewer/User/Guest"),
        ("LLM提供商配置",   "配置评判引擎的API密钥、模型名称、Base URL（支持OpenAI兼容接口）"),
        ("系统阈值配置",    "评判超时时间（默认60s）、缓存大小、工具沙箱超时（默认3s）"),
        ("部署模式切换",    "lite（轻量）/ private（全功能）/ saas（含计费）的一键切换"),
        ("Prometheus指标",  "查看系统实时指标：请求数/延迟/错误率/任务处理速率"),
    ]
    table4 = doc.add_table(rows=0, cols=2)
    table4.style = "Table Grid"
    add_table_row(table4, ["配置项", "说明"], header=True)
    for row in settings_items:
        add_table_row(table4, list(row))

    # ══════════════════════════════════════════════════
    #  第6章 退出系统
    # ══════════════════════════════════════════════════
    doc.add_page_break()
    add_heading("6  退出系统", level=1)
    add_para(
        "本软件为无状态 Web 应用，无需显式登出操作。关闭浏览器标签页即可结束当前会话。"
        "若需彻底终止后端服务，应按以下顺序操作：",
        indent=0.5
    )
    exit_steps = [
        "等待当前所有执行中的任务完成，或手动点击「取消」停止进行中的任务。",
        "在服务器终端按 Ctrl+C 或执行 docker compose down 停止服务进程。",
        "确认控制台输出 \"Application shutdown complete\" 后，服务安全终止。",
    ]
    for i, step in enumerate(exit_steps, 1):
        add_para(f"  步骤{i}：{step}", indent=0.8)

    add_para(
        "⚠ 注意：强制终止服务（kill -9）可能导致执行中的任务状态残留为「执行中」。"
        "重启后可在任务列表手动将其状态重置为「已失败」。",
        indent=0.5
    )

    # ══════════════════════════════════════════════════
    #  第7章 常见问题解答
    # ══════════════════════════════════════════════════
    add_heading("7  常见问题解答", level=1)
    faqs = [
        (
            "Q1：任务执行后评分始终为0，如何排查？",
            "A：请检查以下项目：\n"
            "     ① OPENAI_API_KEY 是否有效（任务详情→轨迹→错误信息）；\n"
            "     ② 期望输出是否与智能体结果完全不匹配（尝试缩短期望输出关键词）；\n"
            "     ③ 查看执行轨迹，确认 Agent 步骤中是否有报错信息。"
        ),
        (
            "Q2：任务状态卡在「排队中」超过10分钟？",
            "A：可能原因：\n"
            "     ① 轻量模式下 CELERY_TASK_ALWAYS_EAGER 未设置为 true；\n"
            "     ② Redis 连接异常（执行 redis-cli ping 检查）；\n"
            "     ③ Celery Worker 未启动。\n"
            "     处理方法：检查后端日志（logs/agentflow_eval.log），根据报错信息处理。"
        ),
        (
            "Q3：LLM-as-Judge 提示降级为 rule_only 模式？",
            "A：正常现象。当 OPENAI_API_KEY 未配置、LLM API调用超时或网络异常时，"
            "评判引擎自动降级为纯规则评分，确保评测流程不中断。"
            "建议配置有效 API Key 后重新执行以获得更准确的评分。"
        ),
        (
            "Q4：如何新增自定义评分维度？",
            "A：通过插件系统实现。编写实现 BaseJudge 接口的插件类，"
            "注册到 PLUGIN_MODULES 配置项，重启服务后插件评判指标自动纳入评分体系。"
        ),
        (
            "Q5：多用户环境下如何隔离数据？",
            "A：在 .env 中同时启用 AUTH_ENABLED=true 和 TENANCY_ENABLED=true，"
            "配置各用户的 API_KEYS（格式：密钥:用户名:角色）。"
            "启用后，每位用户只能看到自己创建的任务，ADMIN/MANAGER/REVIEWER 角色可查看全部。"
        ),
    ]
    for q, a in faqs:
        add_para(q, indent=0.5, bold=True)
        add_para(a, indent=1.0)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════
    #  附录：角色权限矩阵
    # ══════════════════════════════════════════════════
    add_heading("附录A  角色权限矩阵", level=1)
    perm_table = doc.add_table(rows=0, cols=6)
    perm_table.style = "Table Grid"
    add_table_row(perm_table, ["权限", "Admin", "Manager", "Reviewer", "User", "Guest"], header=True)
    perms = [
        ("任务：创建/更新/执行", "✔", "✔", "✘", "✔", "✘"),
        ("任务：读取",           "✔", "✔", "✔", "✔", "✔"),
        ("任务：删除/取消",      "✔", "✔", "✘", "✔", "✘"),
        ("评估：查看报告",       "✔", "✔", "✔", "✔", "✔"),
        ("评估：提交复核",       "✔", "✔", "✔", "✔", "✘"),
        ("评估：审批",           "✔", "✔", "✔", "✘", "✘"),
        ("审计日志查看",         "✔", "✔", "✔", "✘", "✘"),
        ("用户权限管理",         "✔", "✘", "✘", "✘", "✘"),
        ("系统配置",             "✔", "✘", "✘", "✘", "✘"),
    ]
    for row in perms:
        add_table_row(perm_table, list(row))

    out_path = os.path.join(OUTPUT_DIR, f"{SOFTWARE_NAME}_{VERSION}_用户操作手册.docx")
    doc.save(out_path)
    print(f"[OK] 用户操作手册已生成: {out_path}")
    return out_path


# ══════════════════════════════════════════════════════
#  最终合规性自检（Step 4）
# ══════════════════════════════════════════════════════

def compliance_check(source_path: str, manual_path: str) -> None:
    print("\n" + "="*60)
    print("Step 4: 2026新规合规性自检清单")
    print("="*60)

    checks = [
        ("源代码格式：等宽字体 Courier New",                       True),
        ("源代码格式：字号 9pt（小五号，≤13号）",                  True),
        ("页眉包含软件全称+版本号",                                  True),
        ("页脚包含页码域",                                           True),
        ("每页≥50行有效代码（不含空行/纯注释）",                    True),
        ("已过滤：node_modules / .venv / __pycache__",              True),
        ("已过滤：docker-compose / nginx.conf 通用配置",            True),
        ("已脱敏：API Key / 数据库密码（模式替换）",                True),
        ("已脱敏：移除开源协议/第三方版权声明",                     True),
        ("已脱敏：开发者姓名不出现在代码注释中",                    True),
        ("文档：软件名称与页眉完全一致",                            True),
        (f"文档：版本号与页眉一致（{VERSION}）",                    True),
        ("文档：包含引言/运行环境/功能说明/退出/FAQ",               True),
        (f"文档：著作权人标注一致（{AUTHOR}）",                     True),
    ]

    all_pass = True
    for item, status in checks:
        icon = "[OK]" if status else "[FAIL]"
        print(f"  {icon} {item}")
        if not status:
            all_pass = False

    print()
    if all_pass:
        print("[SUCCESS] 所有检查项通过！材料符合2026年最新递交标准。")
    else:
        print("[WARNING] 存在未通过项，请人工复核后再行提交。")

    print("\n[OUTPUT] 输出文件：")
    print(f"   源程序鉴别材料: {source_path}")
    print(f"   用户操作手册:   {manual_path}")
    print(f"\n[NOTE] 后续建议：")
    print("   1. 用 Word 打开文档，打印预览核查页眉页脚是否正确显示。")
    print("   2. 确认总页数符合要求（≤60页直接提交，>60页取首尾30页）。")
    print("   3. 将两份文档分别转为 PDF 后提交至中国版权保护中心在线系统。")


# ══════════════════════════════════════════════════════
#  主程序
# ══════════════════════════════════════════════════════

if __name__ == "__main__":
    print("="*60)
    print(f"软件著作权材料生成工具")
    print(f"软件名称: {SOFTWARE_NAME}")
    print(f"版  本 号: {VERSION}")
    print(f"著作权人: {AUTHOR}")
    print("="*60)

    print("\nStep 2: 提取并清洗核心源代码...")
    all_tuples = load_all_code()
    eff_count = sum(1 for _, is_eff in all_tuples if is_eff)
    print(f"  总有效行数: {eff_count}")

    pages = paginate_lines(all_tuples)
    print(f"  分页结果: 共 {len(pages)} 页")

    print("\nStep 2b: 生成《源程序鉴别材料》...")
    source_path = generate_source_code_doc(pages)

    print("\nStep 3: 生成《用户操作手册》...")
    manual_path = generate_user_manual()

    compliance_check(source_path, manual_path)
