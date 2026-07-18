#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将软著「源程序鉴别材料」导出为 Word（.docx）。

排版规范：
  - A4 竖版
  - 正文 9pt Courier New（中文自动回退系统字体）
  - 固定行距 15 磅；每页 50 行
  - 页眉：软件名称 + 版本号
  - 页脚居中：第 X 页 / 共 Y 页
  - 左侧连续行号

数据来源（优先）：
  1) docs/soft-copyright/generated/源程序鉴别材料_PDF同源.txt
  2) 若无，则调用 generate_source_deposit_pdf 同款流水线生成行数据

用法：
  python scripts/export_source_deposit_docx.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "docs" / "soft-copyright" / "generated"
SRC_TXT = OUT_DIR / "源程序鉴别材料_PDF同源.txt"
DOCX_NAME = "源程序鉴别材料_AgentFlow-Eval_V1.0.docx"

SOFTWARE_NAME = "AgentFlow-Eval Agent自动化评测工作台"
VERSION = "V1.0"
HEADER_TEXT = f"软件名称：{SOFTWARE_NAME} {VERSION}"
LINES_PER_PAGE = 50


def parse_deposit_txt(path: Path) -> list[str]:
    """从 PDF 同源 TXT 解析出正文行（已含行号前缀时去掉前缀）。"""
    lines: list[str] = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        if raw.startswith("========"):
            continue
        m = re.match(r"^\s*(\d+)\|(.*)$", raw)
        if m:
            lines.append(m.group(2))
    return lines


def load_material() -> list[str]:
    if SRC_TXT.exists():
        material = parse_deposit_txt(SRC_TXT)
        if material:
            print(f"  数据源：{SRC_TXT.name}（{len(material)} 行）")
            return material

    # 回退：调用同目录生成器流水线
    print("  未找到 PDF 同源 TXT，尝试重新构建材料…")
    sys.path.insert(0, str(ROOT / "scripts"))
    import generate_source_deposit_pdf as gen  # type: ignore

    files = gen.collect_files()
    stream: list[str] = []
    for fp in files:
        deposit, _effective, _st = gen.process_file(fp)
        if deposit:
            stream.extend(deposit)
    stream = gen.ensure_comment_rate(stream)
    material, _mode, _info = gen.select_head_tail(stream)
    # HEAD+TAIL 在 gen 里可能有自定义路径；若返回 FULL 亦可
    if not material:
        raise SystemExit("无法构建源程序材料")
    return material


def build_docx(material: list[str], out_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    # 补齐为 50 的倍数
    rem = len(material) % LINES_PER_PAGE
    if rem:
        material = material + [""] * (LINES_PER_PAGE - rem)
    total_pages = len(material) // LINES_PER_PAGE

    doc = Document()

    # —— 页面：A4，边距约 1.5cm / 2cm ——
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    # —— 默认样式：9pt 等宽，固定 15 磅行距 ——
    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(9)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(15)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # —— 页眉 ——
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run(HEADER_TEXT)
    hr.font.name = "宋体"
    hr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    hr.font.size = Pt(9)
    hr.bold = True
    # 页眉下边框
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # —— 页脚：第 X 页 / 共 Y 页 ——
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 上边框
    fpPr = fp._p.get_or_add_pPr()
    fBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "666666")
    fBdr.append(top)
    fpPr.append(fBdr)

    def add_run(paragraph, text: str, *, bold: bool = False, size: float = 9, font: str = "宋体"):
        r = paragraph.add_run(text)
        r.font.name = font
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        r.font.size = Pt(size)
        r.bold = bold
        return r

    def add_page_number_field(paragraph):
        """插入 PAGE / NUMPAGES 域：第 X 页 / 共 Y 页。"""
        add_run(paragraph, "第 ", size=9)
        # PAGE field
        run1 = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run1._r.append(fld_char_begin)

        run2 = paragraph.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run2._r.append(instr)

        run3 = paragraph.add_run()
        fld_char_sep = OxmlElement("w:fldChar")
        fld_char_sep.set(qn("w:fldCharType"), "separate")
        run3._r.append(fld_char_sep)

        run4 = paragraph.add_run("1")
        run4.font.size = Pt(9)
        run4.font.name = "宋体"

        run5 = paragraph.add_run()
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run5._r.append(fld_char_end)

        add_run(paragraph, " 页 / 共 ", size=9)

        # NUMPAGES field
        run6 = paragraph.add_run()
        b2 = OxmlElement("w:fldChar")
        b2.set(qn("w:fldCharType"), "begin")
        run6._r.append(b2)

        run7 = paragraph.add_run()
        instr2 = OxmlElement("w:instrText")
        instr2.set(qn("xml:space"), "preserve")
        instr2.text = " NUMPAGES "
        run7._r.append(instr2)

        run8 = paragraph.add_run()
        s2 = OxmlElement("w:fldChar")
        s2.set(qn("w:fldCharType"), "separate")
        run8._r.append(s2)

        run9 = paragraph.add_run(str(total_pages))
        run9.font.size = Pt(9)
        run9.font.name = "宋体"

        run10 = paragraph.add_run()
        e2 = OxmlElement("w:fldChar")
        e2.set(qn("w:fldCharType"), "end")
        run10._r.append(e2)

        add_run(paragraph, " 页", size=9)

    add_page_number_field(fp)

    # 制表位：行号列宽约 1.1cm
    tab_pos = Cm(1.1)

    gno = 1
    for page_idx in range(total_pages):
        chunk = material[page_idx * LINES_PER_PAGE : (page_idx + 1) * LINES_PER_PAGE]
        if page_idx > 0:
            # 分页
            doc.add_page_break()

        for text in chunk:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(15)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.LEFT)

            # 行号（灰色）+ Tab + 正文
            display = text if len(text) <= 100 else text[:97] + "..."
            r_no = p.add_run(f"{gno:>5d}")
            r_no.font.name = "Courier New"
            r_no._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
            r_no._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
            r_no._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r_no.font.size = Pt(8)
            r_no.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            r_tab = p.add_run("\t")
            r_tab.font.size = Pt(9)

            r_body = p.add_run(display)
            r_body.font.name = "Courier New"
            r_body._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
            r_body._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
            r_body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r_body.font.size = Pt(9)

            gno += 1

    # 核心属性
    core = doc.core_properties
    core.title = f"{SOFTWARE_NAME} {VERSION} 源程序鉴别材料"
    core.author = "李凯昕"
    core.subject = "计算机软件著作权登记-源程序鉴别材料"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"  页数（逻辑）：{total_pages}")
    print(f"  输出：{out_path}")


def main() -> int:
    print("=" * 60)
    print("导出软著源程序鉴别材料 → Word (.docx)")
    print(f"  {HEADER_TEXT}")
    print("=" * 60)

    try:
        from docx import Document  # noqa: F401
    except ImportError:
        print("缺少 python-docx，请执行：pip install python-docx", file=sys.stderr)
        return 1

    material = load_material()
    if len(material) < LINES_PER_PAGE:
        print(f"材料行数过少：{len(material)}", file=sys.stderr)
        return 1

    out = OUT_DIR / DOCX_NAME
    build_docx(material, out)
    size_kb = out.stat().st_size / 1024
    print(f"  大小：{size_kb:.1f} KB")
    print("完成。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
